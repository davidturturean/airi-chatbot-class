#!/usr/bin/env python3
"""
Create an enhanced UI + Performance feedback document with performance elements highlighted.
"""

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_COLOR_INDEX
except ImportError:
    import sys
    import subprocess
    print("python-docx not found. Installing...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_COLOR_INDEX

def add_highlighted_text(paragraph, text):
    """Add white text with black background (inverted highlighting)."""
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor(255, 255, 255)  # White text
    run.font.highlight_color = WD_COLOR_INDEX.BLACK  # Black highlight
    return run

def create_ui_performance_document():
    """Create the enhanced UI + Performance feedback document."""
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = "AI Risk Repository Chatbot: User Feedback & Performance Improvement Plan"
    doc.core_properties.author = "MIT FutureTech"
    
    # Title
    title = doc.add_heading('AI Risk Repository Chatbot: User Feedback & ', 0)
    add_highlighted_text(title, 'Performance ')
    title.add_run('Improvement Plan')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Opening paragraph
    p = doc.add_paragraph(
        'The AI Risk Repository Chatbot has been integrated as a proof-of-concept into the airisk.mit.edu '
        'mirror via Webflow. To optimize user experience '
    )
    add_highlighted_text(p, 'and system performance')
    p.add_run(', ensuring the tool meets researcher and practitioner needs, I propose a structured feedback ')
    add_highlighted_text(p, 'and performance testing ')
    p.add_run('gathering initiative within the FutureTech team and adjacent stakeholders. Let this be the document '
              'outlining the methodology for collecting actionable insights')
    add_highlighted_text(p, ' and performance metrics')
    p.add_run('.')
    
    # Current State & Challenge
    doc.add_heading('Current State & Challenge', 1)
    
    doc.add_heading('What We Have', 2)
    doc.add_paragraph('• Chatbot embedded in airisk.mit.edu')
    doc.add_paragraph('• Direct navigation from main site "Chatbot" tab to /chat interface (no separate homepage for the chatbot)')
    doc.add_paragraph('• Citation system linking to source documents')
    p = doc.add_paragraph('• Diverse query processing ')
    add_highlighted_text(p, 'averaging 2-3 second response times')
    
    doc.add_heading('The Challenge', 2)
    p = doc.add_paragraph(
        'The "curse of knowledge", of understanding the system\'s capabilities and limitations. '
        'We need fresh perspectives to understand:'
    )
    doc.add_paragraph('• What new users expect vs. what they experience')
    doc.add_paragraph('• Whether the current interface supports user goals effectively')
    doc.add_paragraph('• What barriers prevent successful task completion')
    p = doc.add_paragraph('• ')
    add_highlighted_text(p, 'Which queries cause performance bottlenecks')
    p = doc.add_paragraph('• ')
    add_highlighted_text(p, 'How system accuracy varies by query type')
    
    # Key Questions
    doc.add_heading('Key Questions to Answer Internally', 1)
    
    doc.add_heading('1. Onboarding & First Impressions', 2)
    doc.add_paragraph('• Do users understand what the chatbot can/cannot do from first impression?')
    doc.add_paragraph('• Are example queries or guided options needed?')
    p = doc.add_paragraph('• ')
    add_highlighted_text(p, 'Does initial query response time affect user engagement?')
    
    doc.add_heading('2. Interface Design Decisions (not mutually exclusive)', 2)
    doc.add_paragraph('Current: Full-page embedded chat via iframe')
    doc.add_paragraph('Alternative A: Floating chat widget that expands')
    p = doc.add_paragraph(
        '→ Idea: implement Alternative A as well, with option to expand to full-page, and see how the testing for that goes'
    )
    p = doc.add_paragraph('→ ')
    add_highlighted_text(p, 'Measure load time differences between implementations')
    
    doc.add_heading('3. User Intent & Success', 2)
    doc.add_paragraph('• What are users actually trying to accomplish?')
    doc.add_paragraph('• Can they complete their intended tasks?')
    doc.add_paragraph('• What queries fail most often?')
    p = doc.add_paragraph('• ')
    add_highlighted_text(p, 'What is the accuracy rate for different query categories?')
    p = doc.add_paragraph('• ')
    add_highlighted_text(p, 'How many retries occur due to unsatisfactory responses?')
    
    doc.add_heading('4. Trust & Understanding', 2)
    doc.add_paragraph('• Do users trust the responses?')
    doc.add_paragraph('• Do users understand the citations and use them?')
    p = doc.add_paragraph('• ')
    add_highlighted_text(p, 'Are citations accurate and do links work properly?')
    p = doc.add_paragraph('• ')
    add_highlighted_text(p, 'What percentage of responses contain verifiable facts?')
    
    # Feedback Methodology
    doc.add_heading('Feedback Methodology', 1)
    
    doc.add_heading('Phase 1: User Research (Weeks 1-2?)', 2)
    
    doc.add_heading('Structured User Testing Sessions', 3)
    p = doc.add_paragraph('Aim for 6 people? ')
    add_highlighted_text(p, 'Plus automated testing on 100+ queries')
    
    doc.add_paragraph()
    doc.add_paragraph('Specific test tasks to fulfill using the chatbot:')
    doc.add_paragraph('1. Find information about AI privacy risks')
    doc.add_paragraph('2. Understand how the repository categorizes risks')
    doc.add_paragraph('3. Explore employment impacts of AI')
    doc.add_paragraph('4. Locate specific statistics about AI safety')
    
    p = doc.add_paragraph()
    add_highlighted_text(p, 'Performance metrics for each task:')
    p = doc.add_paragraph()
    add_highlighted_text(p, '• Response time (target <2s)')
    p = doc.add_paragraph()
    add_highlighted_text(p, '• Factual accuracy (target >95%)')
    p = doc.add_paragraph()
    add_highlighted_text(p, '• Citation validity (100% working links)')
    
    doc.add_paragraph()
    doc.add_paragraph('Non-specific test tasks:')
    doc.add_paragraph('1. 4-6 prompts/tasks at the liberty of the tester (with eventual short follow-ups)')
    
    doc.add_paragraph()
    doc.add_paragraph('What we will measure:')
    doc.add_paragraph('• Task completion rates')
    doc.add_paragraph('• Understanding of system capabilities')
    doc.add_paragraph('• Points of confusion or abandonment')
    doc.add_paragraph('• Time to first meaningful interaction? (first meaningful question asked and properly responded)')
    doc.add_paragraph('• User satisfaction with the responses')
    p = doc.add_paragraph('• ')
    add_highlighted_text(p, 'Query-to-response latency for each interaction')
    p = doc.add_paragraph('• ')
    add_highlighted_text(p, 'Token usage per query type')
    p = doc.add_paragraph('• ')
    add_highlighted_text(p, 'Error rates and timeout frequencies')
    
    doc.add_paragraph()
    doc.add_paragraph('How we will measure:')
    doc.add_paragraph('Deploying survey after users interact with the chatbot:')
    doc.add_paragraph('1. Did you find what you were looking for? (Yes/Partially/No), for each question posed to the chatbot')
    doc.add_paragraph('2. What was confusing about the interface? (open text)')
    doc.add_paragraph('3. Can we better convey what the system knows/what the system can do?')
    doc.add_paragraph('4. What would make this tool more useful? (open text)')
    doc.add_paragraph('5. How likely are you to recommend this to someone of a similar occupation to you? '
                      'What about to a family member or friend of a totally different occupation? (1-10 scale)')
    p = doc.add_paragraph('6. ')
    add_highlighted_text(p, 'Was the response speed acceptable? (Yes/No/Sometimes)')
    p = doc.add_paragraph('7. ')
    add_highlighted_text(p, 'Did you encounter any errors or timeouts? (Yes/No + details)')
    
    doc.add_paragraph()
    doc.add_paragraph('+ Reporting the distribution/statistics of answers to numerical scale question')
    p = doc.add_paragraph('+ ')
    add_highlighted_text(p, 'Automated performance metrics dashboard with real-time monitoring')
    
    doc.add_paragraph()
    doc.add_paragraph('+ IF we videotape: empirical measurement of understanding of system capabilities '
                      '& measurement of time to first meaningful interaction')
    p = doc.add_paragraph('+ ')
    add_highlighted_text(p, 'Screen recordings to correlate UI actions with performance spikes')
    
    doc.add_heading('Secondary for now, important in the long run', 3)
    doc.add_paragraph('1. How well does the iframe embed work on phones/tablets?')
    p = doc.add_paragraph('   ')
    add_highlighted_text(p, '→ Mobile performance benchmarks')
    doc.add_paragraph('2. Should we add autocomplete?')
    p = doc.add_paragraph('   ')
    add_highlighted_text(p, '→ Impact on query formulation time')
    doc.add_paragraph('3. Should we add conversation history/export?')
    p = doc.add_paragraph('   ')
    add_highlighted_text(p, '→ Session memory usage implications')
    
    p = doc.add_paragraph()
    add_highlighted_text(p, 'Additional Performance Testing:')
    p = doc.add_paragraph()
    add_highlighted_text(p, '• Load testing with 10, 50, 100 concurrent users')
    p = doc.add_paragraph()
    add_highlighted_text(p, '• Stress testing with edge cases and malformed queries')
    p = doc.add_paragraph()
    add_highlighted_text(p, '• Database query optimization opportunities')
    p = doc.add_paragraph()
    add_highlighted_text(p, '• Cache effectiveness measurement')
    
    # Phase 2
    doc.add_heading('Phase 2: Analysis and Implementation (Week 3-4?)', 2)
    doc.add_paragraph('• Synthesizing findings into priority recommendations')
    p = doc.add_paragraph('   ')
    add_highlighted_text(p, '→ Performance bottleneck identification')
    doc.add_paragraph('• And implement priority features')
    doc.add_paragraph('• Fix identified pain points')
    p = doc.add_paragraph('   ')
    add_highlighted_text(p, '→ Optimize slow query patterns')
    doc.add_paragraph('• Implement high-priority features')
    p = doc.add_paragraph('   ')
    add_highlighted_text(p, '→ Deploy performance improvements')
    
    p = doc.add_paragraph()
    p = doc.add_paragraph()
    add_highlighted_text(p, 'Performance Success Criteria:')
    p = doc.add_paragraph()
    add_highlighted_text(p, '• Response time <2s (P95)')
    p = doc.add_paragraph()
    add_highlighted_text(p, '• Accuracy >95% for repository facts')
    p = doc.add_paragraph()
    add_highlighted_text(p, '• Error rate <2%')
    p = doc.add_paragraph()
    add_highlighted_text(p, '• All citations functional')
    p = doc.add_paragraph()
    add_highlighted_text(p, '• Support 100+ concurrent users')
    
    # Save document
    doc.save('UI_AND_PERFORMANCE_FEEDBACK_PLAN.docx')
    print('Successfully created: UI_AND_PERFORMANCE_FEEDBACK_PLAN.docx')
    print('Performance elements are highlighted with white text on black background')


if __name__ == '__main__':
    create_ui_performance_document()