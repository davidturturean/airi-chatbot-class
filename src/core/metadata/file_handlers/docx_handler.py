"""
DOCX (Word document) file handler.
"""
from pathlib import Path
from typing import Dict, List, Any, Optional
from .base_handler import BaseFileHandler
from ....config.logging import get_logger

logger = get_logger(__name__)

class DocxHandler(BaseFileHandler):
    """Handler for Microsoft Word documents (.docx)."""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.docx']
    
    def can_handle(self, file_path: Path) -> bool:
        """Check if this is a DOCX file."""
        return file_path.suffix.lower() in self.supported_extensions
    
    def extract_data(self, file_path: Path) -> List[Dict[str, Any]]:
        """Extract structured data from DOCX files."""
        try:
            # Try to import python-docx
            try:
                from docx import Document
            except ImportError:
                logger.error("python-docx not installed. Install with: pip install python-docx")
                # Fallback: treat as binary and extract basic info
                return self._extract_basic_info(file_path)
            
            # Read the document
            doc = Document(file_path)
            extracted_data = []
            
            # Extract paragraphs
            paragraphs = self._extract_paragraphs(doc)
            if paragraphs:
                extracted_data.append(paragraphs)
            
            # Extract tables
            tables = self._extract_tables(doc)
            extracted_data.extend(tables)
            
            # Extract lists
            lists = self._extract_lists(doc)
            if lists:
                extracted_data.append(lists)
            
            # Extract document properties
            properties = self._extract_properties(doc, file_path)
            if properties:
                extracted_data.append(properties)
            
            return extracted_data
            
        except Exception as e:
            logger.error(f"Error extracting data from DOCX {file_path}: {str(e)}")
            return self._extract_basic_info(file_path)
    
    def _extract_paragraphs(self, doc) -> Optional[Dict[str, Any]]:
        """Extract paragraphs from document."""
        paragraphs = []
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                paragraph_data = {
                    "paragraph_number": i + 1,
                    "content": text,
                    "style": para.style.name if para.style else "Normal",
                    "length": len(text)
                }
                
                # Check if it's a heading
                if para.style and 'Heading' in para.style.name:
                    paragraph_data["is_heading"] = True
                    paragraph_data["heading_level"] = self._extract_heading_level(para.style.name)
                
                paragraphs.append(paragraph_data)
        
        if paragraphs:
            return {
                "table_name": self.sanitize_table_name("paragraphs"),
                "data": paragraphs,
                "metadata": {
                    "extraction_method": "docx_paragraphs",
                    "paragraph_count": len(paragraphs)
                }
            }
        
        return None
    
    def _extract_tables(self, doc) -> List[Dict[str, Any]]:
        """Extract tables from document."""
        tables = []
        
        for table_idx, table in enumerate(doc.tables):
            # Extract table data
            table_data = []
            headers = None
            
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                
                if row_idx == 0 and all(cell for cell in row_data):
                    # First row might be headers
                    headers = [self._clean_column_name(h) for h in row_data]
                else:
                    if headers and len(row_data) == len(headers):
                        # Create dict with headers
                        table_data.append(dict(zip(headers, row_data)))
                    else:
                        # Create dict with generic column names
                        row_dict = {}
                        for col_idx, value in enumerate(row_data):
                            row_dict[f"column_{col_idx + 1}"] = value
                        table_data.append(row_dict)
            
            if table_data:
                tables.append({
                    "table_name": self.sanitize_table_name(f"table_{table_idx + 1}"),
                    "data": table_data,
                    "metadata": {
                        "extraction_method": "docx_table",
                        "table_index": table_idx + 1,
                        "row_count": len(table.rows),
                        "column_count": len(table.columns)
                    }
                })
        
        return tables
    
    def _extract_lists(self, doc) -> Optional[Dict[str, Any]]:
        """Extract bulleted/numbered lists from document."""
        lists = []
        current_list = []
        list_type = None
        
        for para in doc.paragraphs:
            # Check if paragraph is part of a list
            if para.style and ('List' in para.style.name or 
                              para.text.strip().startswith(('•', '-', '*', '1.', '2.', '3.'))):
                
                # Determine list type
                if para.text.strip()[0].isdigit():
                    current_type = 'numbered'
                else:
                    current_type = 'bulleted'
                
                # If list type changes, save current list
                if list_type and list_type != current_type and current_list:
                    lists.append({
                        "list_type": list_type,
                        "items": current_list,
                        "item_count": len(current_list)
                    })
                    current_list = []
                
                list_type = current_type
                
                # Clean and add item
                text = para.text.strip()
                # Remove list markers
                text = text.lstrip('•-*').strip()
                text = text.lstrip('0123456789.').strip()
                
                if text:
                    current_list.append(text)
            
            elif current_list:
                # End of list
                lists.append({
                    "list_type": list_type,
                    "items": current_list,
                    "item_count": len(current_list)
                })
                current_list = []
                list_type = None
        
        # Save last list if any
        if current_list:
            lists.append({
                "list_type": list_type,
                "items": current_list,
                "item_count": len(current_list)
            })
        
        if lists:
            # Flatten lists into records
            list_items = []
            for list_idx, lst in enumerate(lists):
                for item_idx, item in enumerate(lst['items']):
                    list_items.append({
                        "list_number": list_idx + 1,
                        "list_type": lst['list_type'],
                        "item_number": item_idx + 1,
                        "content": item
                    })
            
            return {
                "table_name": self.sanitize_table_name("lists"),
                "data": list_items,
                "metadata": {
                    "extraction_method": "docx_lists",
                    "list_count": len(lists),
                    "total_items": len(list_items)
                }
            }
        
        return None
    
    def _extract_properties(self, doc, file_path: Path) -> Optional[Dict[str, Any]]:
        """Extract document properties."""
        properties = []
        
        try:
            core_props = doc.core_properties
            
            # Extract available properties
            prop_data = {
                "property": "title",
                "value": core_props.title or "N/A"
            }
            properties.append(prop_data)
            
            if core_props.author:
                properties.append({
                    "property": "author",
                    "value": core_props.author
                })
            
            if core_props.created:
                properties.append({
                    "property": "created",
                    "value": str(core_props.created)
                })
            
            if core_props.modified:
                properties.append({
                    "property": "modified",
                    "value": str(core_props.modified)
                })
            
            # Add file info
            properties.append({
                "property": "file_name",
                "value": file_path.name
            })
            
            properties.append({
                "property": "file_size",
                "value": str(file_path.stat().st_size)
            })
            
        except Exception as e:
            logger.warning(f"Could not extract properties: {str(e)}")
        
        if properties:
            return {
                "table_name": self.sanitize_table_name("document_properties"),
                "data": properties,
                "metadata": {
                    "extraction_method": "docx_properties"
                }
            }
        
        return None
    
    def _extract_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name."""
        import re
        match = re.search(r'Heading (\d+)', style_name)
        if match:
            return int(match.group(1))
        return 0
    
    def _extract_basic_info(self, file_path: Path) -> List[Dict[str, Any]]:
        """Fallback: extract basic file information."""
        try:
            stat = file_path.stat()
            return [{
                "table_name": self.sanitize_table_name("file_info"),
                "data": [{
                    "file_name": file_path.name,
                    "file_size": stat.st_size,
                    "file_type": "docx",
                    "error": "python-docx not installed for full extraction"
                }],
                "metadata": {
                    "extraction_method": "basic_file_info"
                }
            }]
        except Exception as e:
            logger.error(f"Could not extract basic info: {str(e)}")
            return []
    
    def _clean_column_name(self, name: str) -> str:
        """Clean column name for SQL compatibility."""
        import re
        name = re.sub(r'[^a-zA-Z0-9_\s]', '', name)
        name = name.replace(' ', '_')
        name = re.sub(r'_+', '_', name)
        name = name.strip('_').lower()
        
        if not name:
            return "unnamed"
        if not name[0].isalpha():
            name = f"col_{name}"
        
        return name