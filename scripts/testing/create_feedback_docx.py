#!/usr/bin/env python3
"""
Convert the feedback plan markdown to a Word document.
"""

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    import sys
    import subprocess
    print("python-docx not found. Installing...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
import re

def create_feedback_document():
    # Create a new document
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = "AI Risk Repository Chatbot: User Feedback & Improvement Plan"
    doc.core_properties.author = "MIT FutureTech"
    
    # Add title
    title = doc.add_heading('AI Risk Repository Chatbot: User Feedback & Improvement Plan', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Executive Summary
    doc.add_heading('Executive Summary', 1)
    doc.add_paragraph(
        'The AI Risk Repository Chatbot has been successfully integrated into airisk.mit.edu via Webflow embedding. '
        'To optimize user experience and ensure the tool meets researcher and practitioner needs, we propose a structured '
        'feedback gathering initiative with the FutureTech team and adjacent stakeholders. This document outlines our '
        'methodology for collecting actionable insights and implementing data-driven improvements.'
    )
    
    # Current State & Challenge
    doc.add_heading('Current State & Challenge', 1)
    
    doc.add_heading('What We Have', 2)
    items = [
        'Fully functional chatbot embedded in airisk.mit.edu',
        'Direct navigation from main site "Chatbot" tab to /chat interface',
        'Working query processing for 1,612 AI risks across 7 domains',
        'Citation system linking to source documents'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('The Challenge', 2)
    p = doc.add_paragraph(
        'As developers, we have "curse of knowledge" - we understand the system\'s capabilities and limitations intimately. '
        'We need fresh perspectives to understand:'
    )
    challenges = [
        'What new users expect vs. what they experience',
        'Whether the current interface supports user goals effectively',
        'What barriers prevent successful task completion'
    ]
    for challenge in challenges:
        doc.add_paragraph(challenge, style='List Bullet')
    
    # Proposed Feedback Methodology
    doc.add_heading('Proposed Feedback Methodology', 1)
    
    doc.add_heading('Phase 1: User Research (Weeks 1-2)', 2)
    
    doc.add_heading('A. Structured User Testing Sessions', 3)
    
    p = doc.add_paragraph()
    p.add_run('Participants: ').bold = True
    p.add_run('10-12 people (45 minutes each)')
    
    participants = [
        '4 FutureTech team members',
        '3 AI/ML researchers (external to MIT)',
        '3 policy/governance professionals',
        '2 graduate students in relevant fields'
    ]
    for participant in participants:
        doc.add_paragraph(participant, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('Test Tasks:').bold = True
    tasks = [
        'Find information about AI privacy risks',
        'Understand how the repository categorizes risks',
        'Explore employment impacts of AI',
        'Locate specific statistics about AI safety'
    ]
    for i, task in enumerate(tasks, 1):
        doc.add_paragraph(f'{i}. {task}', style='List Number')
    
    p = doc.add_paragraph()
    p.add_run('What We\'ll Measure:').bold = True
    measures = [
        'Task completion rates',
        'Time to first meaningful interaction',
        'Points of confusion or abandonment',
        'Understanding of system capabilities'
    ]
    for measure in measures:
        doc.add_paragraph(measure, style='List Bullet')
    
    doc.add_heading('B. Quick Intercept Surveys', 3)
    doc.add_paragraph('Deploy a brief 3-question survey after users interact with the chatbot:')
    survey_questions = [
        'Did you find what you were looking for? (Yes/Partially/No)',
        'What would make this tool more useful? (open text)',
        'How likely are you to recommend this? (1-10 scale)'
    ]
    for i, q in enumerate(survey_questions, 1):
        doc.add_paragraph(f'{i}. {q}', style='List Number')
    
    doc.add_heading('Phase 2: Analysis & Prioritization (Week 3)', 2)
    phase2_items = [
        'Synthesize findings into key themes',
        'Create priority matrix (High/Low Impact vs. Easy/Hard Implementation)',
        'Develop specific design recommendations'
    ]
    for item in phase2_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Phase 3: Implementation (Weeks 4-6)', 2)
    phase3_items = [
        'Address quick wins immediately',
        'Design solutions for major issues',
        'A/B test significant changes'
    ]
    for item in phase3_items:
        doc.add_paragraph(item, style='List Bullet')
    
    # Key Questions to Answer
    doc.add_heading('Key Questions to Answer', 1)
    
    doc.add_heading('1. Onboarding & First Impressions', 2)
    onboarding_questions = [
        'Should users see a welcome page before the chat interface?',
        'Do users understand what the chatbot can/cannot do?',
        'Are example queries or guided options needed?'
    ]
    for q in onboarding_questions:
        doc.add_paragraph(q, style='List Bullet')
    
    doc.add_heading('2. Interface Design Decisions', 2)
    design_options = [
        ('Current:', 'Full-page embedded chat via iframe'),
        ('Alternative A:', 'Floating chat widget that expands'),
        ('Alternative B:', 'Split interface with browse + chat options'),
        ('Alternative C:', 'Progressive disclosure starting with search bar')
    ]
    for label, desc in design_options:
        p = doc.add_paragraph()
        p.add_run(label).bold = True
        p.add_run(f' {desc}')
    
    doc.add_heading('3. User Intent & Success', 2)
    intent_questions = [
        'What are users actually trying to accomplish?',
        'Can they complete their intended tasks?',
        'What queries fail most often?'
    ]
    for q in intent_questions:
        doc.add_paragraph(q, style='List Bullet')
    
    doc.add_heading('4. Trust & Understanding', 2)
    trust_questions = [
        'Do users trust the responses?',
        'Is the MIT affiliation/credibility clear?',
        'Do users understand citations and use them?'
    ]
    for q in trust_questions:
        doc.add_paragraph(q, style='List Bullet')
    
    # Specific Feedback Areas
    doc.add_heading('Specific Feedback Areas', 1)
    
    doc.add_heading('High Priority', 2)
    high_priority = [
        ('Entry Point:', 'Is direct navigation to /chat optimal, or do users need context?'),
        ('Capability Communication:', 'How can we better convey what the system knows?'),
        ('Error Handling:', 'What happens when users ask out-of-scope questions?'),
        ('Mobile Experience:', 'How well does the iframe embed work on phones/tablets?')
    ]
    for i, (label, desc) in enumerate(high_priority, 1):
        p = doc.add_paragraph(f'{i}. ')
        p.add_run(label).bold = True
        p.add_run(f' {desc}')
    
    doc.add_heading('Medium Priority', 2)
    medium_priority = [
        'Query suggestions or autocomplete',
        'Conversation history/export features',
        'Integration with main repository browse function',
        'Visual indicators of processing/loading states'
    ]
    for i, item in enumerate(medium_priority, 1):
        doc.add_paragraph(f'{i}. {item}')
    
    doc.add_heading('Future Considerations', 2)
    future_items = [
        'Personalization for different user types',
        'Multi-language support expansion',
        'API access for programmatic queries'
    ]
    for i, item in enumerate(future_items, 1):
        doc.add_paragraph(f'{i}. {item}')
    
    # Timeline & Resources
    doc.add_heading('Timeline & Resources', 1)
    
    timeline_sections = [
        ('Week 1-2: Data Collection', [
            'Recruit participants',
            'Conduct user sessions',
            'Deploy analytics and surveys'
        ]),
        ('Week 3: Analysis', [
            'Synthesize findings',
            'Create priority recommendations',
            'Design mockups for major changes'
        ]),
        ('Week 4-6: Implementation', [
            'Fix identified pain points',
            'Implement high-priority features',
            'Deploy A/B tests'
        ])
    ]
    
    for section_title, items in timeline_sections:
        doc.add_heading(section_title, 2)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Required Resources', 2)
    resources = [
        '10-12 hours of user testing time',
        '1-2 team members to conduct sessions',
        'Development time for implementing changes'
    ]
    for resource in resources:
        doc.add_paragraph(resource, style='List Bullet')
    
    # Success Metrics & Expected Outcomes
    doc.add_heading('Success Metrics & Expected Outcomes', 1)
    
    doc.add_heading('Key Performance Indicators', 2)
    kpis = [
        ('Engagement Rate:', '>60% of visitors ask at least one question'),
        ('Success Rate:', '>75% of queries answered satisfactorily'),
        ('Return Rate:', '>30% return within one week'),
        ('Net Promoter Score:', '>7/10 average')
    ]
    for i, (label, target) in enumerate(kpis, 1):
        p = doc.add_paragraph(f'{i}. ')
        p.add_run(label).bold = True
        p.add_run(f' {target}')
    
    doc.add_heading('Deliverables', 2)
    deliverables = [
        ('Feedback Report', 'with synthesized findings and recommendations'),
        ('Priority Roadmap', 'for UI/UX improvements'),
        ('Updated Interface', 'based on user feedback'),
        ('Best Practices Guide', 'for future iterations')
    ]
    for i, (title, desc) in enumerate(deliverables, 1):
        p = doc.add_paragraph(f'{i}. ')
        p.add_run(title).bold = True
        p.add_run(f' {desc}')
    
    # Next Steps
    doc.add_heading('Next Steps', 1)
    next_steps = [
        ('Approval', 'of this feedback plan'),
        ('Recruitment', 'of test participants from FutureTech and adjacent teams'),
        ('Scheduling', 'of user testing sessions'),
        ('Implementation', 'of quick analytics to establish baselines')
    ]
    for i, (action, detail) in enumerate(next_steps, 1):
        p = doc.add_paragraph(f'{i}. ')
        p.add_run(action).bold = True
        p.add_run(f' {detail}')
    
    # Questions for Discussion
    doc.add_heading('Questions for Discussion', 1)
    discussion_questions = [
        'Are there specific user groups we should prioritize?',
        'What aspects of the current implementation are non-negotiable?',
        'What is our budget/timeline for implementing changes?',
        'Should we coordinate with the main airisk.mit.edu redesign efforts?'
    ]
    for i, q in enumerate(discussion_questions, 1):
        doc.add_paragraph(f'{i}. {q}')
    
    # Footer
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Contact: David Turturean | Last Updated: November 2024').italic = True
    
    # Save the document
    doc.save('CHATBOT_UI_FEEDBACK_PLAN.docx')
    print('Document created successfully: CHATBOT_UI_FEEDBACK_PLAN.docx')

if __name__ == '__main__':
    create_feedback_document()