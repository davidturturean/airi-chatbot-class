#!/usr/bin/env python3
"""
Create two comprehensive performance evaluation documents for the AI Risk Repository Chatbot.
"""

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    import sys
    import subprocess
    print("python-docx not found. Installing...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE

def create_developer_testing_document():
    """Create comprehensive developer performance testing document."""
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = "Developer Performance Testing Plan"
    doc.core_properties.author = "MIT FutureTech"
    
    # Title
    title = doc.add_heading('AI Risk Repository Chatbot: Developer Performance Testing Plan', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Executive Summary
    doc.add_heading('Executive Summary', 1)
    doc.add_paragraph(
        'This document provides a comprehensive technical framework for evaluating the performance '
        'of the AI Risk Repository Chatbot. It includes automated testing protocols, performance '
        'metrics collection, system stress testing, and continuous monitoring strategies.'
    )
    
    # 1. Performance Testing Framework
    doc.add_heading('1. Performance Testing Framework', 1)
    
    doc.add_heading('1.1 Testing Architecture', 2)
    doc.add_paragraph(
        'Our performance testing framework consists of multiple layers designed to evaluate '
        'different aspects of system performance:'
    )
    
    layers = [
        'Unit Testing: Individual component performance',
        'Integration Testing: API endpoint response times',
        'System Testing: End-to-end query processing',
        'Load Testing: Concurrent user simulation',
        'Stress Testing: Breaking point identification'
    ]
    for layer in layers:
        doc.add_paragraph(layer, style='List Bullet')
    
    doc.add_heading('1.2 Key Performance Indicators', 2)
    
    # Add KPI table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Target'
    hdr_cells[2].text = 'Acceptable'
    hdr_cells[3].text = 'Critical'
    
    kpis = [
        ('Response Time (P95)', '<2s', '<3s', '>5s'),
        ('Accuracy Rate', '>95%', '>85%', '<85%'),
        ('Error Rate', '<2%', '<5%', '>5%'),
        ('Uptime', '>99.9%', '>99%', '<99%'),
        ('Cache Hit Rate', '>80%', '>60%', '<60%'),
        ('Token Efficiency', '<2000/query', '<3000/query', '>3000/query')
    ]
    
    for kpi in kpis:
        row_cells = table.add_row().cells
        for i, value in enumerate(kpi):
            row_cells[i].text = value
    
    # 2. Ground Truth Dataset
    doc.add_heading('2. Ground Truth Dataset Creation', 1)
    
    doc.add_heading('2.1 Query Categories', 2)
    doc.add_paragraph('Test queries are organized into five categories to ensure comprehensive coverage:')
    
    doc.add_heading('Category A: Factual Queries', 3)
    doc.add_paragraph('Queries with verifiable, exact answers from the repository:')
    
    # Add code block style
    code_para = doc.add_paragraph()
    code_para.add_run('''test_queries_factual = [
    {"query": "How many risks are in the repository?", 
     "expected": "1612", 
     "type": "exact_match",
     "priority": "critical"},
    {"query": "What percentage of risks are in the discrimination domain?", 
     "expected": "15%", 
     "type": "exact_match",
     "priority": "critical"},
    {"query": "How many domains exist?", 
     "expected": "7", 
     "type": "exact_match",
     "priority": "critical"},
    {"query": "What is the largest risk domain?", 
     "expected": "socioeconomic and environmental harms", 
     "type": "contains",
     "priority": "high"}
]''').font.name = 'Courier New'
    
    doc.add_heading('Category B: Taxonomy Queries', 3)
    doc.add_paragraph('Queries testing understanding of risk categorization:')
    
    code_para = doc.add_paragraph()
    code_para.add_run('''test_queries_taxonomy = [
    {"query": "What are the 7 domains of AI risk?", 
     "expected_terms": ["discrimination", "privacy", "misinformation", 
                        "malicious use", "security", 
                        "human-computer interaction", 
                        "socioeconomic and environmental harms"],
     "type": "contains_all",
     "priority": "critical"},
    {"query": "How are risks categorized by entity?", 
     "expected_terms": ["human", "AI", "human & AI"],
     "type": "contains_any",
     "priority": "high"}
]''').font.name = 'Courier New'
    
    doc.add_heading('Category C: Metadata Queries', 3)
    doc.add_paragraph('Queries testing metadata service functionality:')
    
    code_para = doc.add_paragraph()
    code_para.add_run('''test_queries_metadata = [
    {"query": "What are the main risk domains in the AI risk repository?",
     "should_not_error": True,
     "expected_service": "metadata",
     "priority": "critical"},
    {"query": "Show me risks by timing",
     "expected_terms": ["pre-deployment", "post-deployment"],
     "type": "classification",
     "priority": "high"}
]''').font.name = 'Courier New'
    
    doc.add_heading('Category D: Edge Cases', 3)
    doc.add_paragraph('Queries testing system robustness and error handling:')
    
    edge_cases = [
        'Empty query: ""',
        'Single character: "a"',
        'Special characters: "!@#$%^&*()"',
        'SQL injection attempt: "; DROP TABLE risks;--"',
        'Extremely long query: 5000+ characters',
        'Non-English characters: "风险是什么？"',
        'Mixed languages: "What are the risques principales?"'
    ]
    for case in edge_cases:
        doc.add_paragraph(case, style='List Bullet')
    
    doc.add_heading('Category E: Performance Stress Queries', 3)
    doc.add_paragraph('Queries designed to stress specific components:')
    
    stress_queries = [
        'Complex multi-domain comparison queries',
        'Queries requiring extensive document retrieval',
        'Queries with high token generation requirements',
        'Rapid sequential queries (rate limiting test)',
        'Queries requiring multiple service calls'
    ]
    for query in stress_queries:
        doc.add_paragraph(query, style='List Bullet')
    
    # 3. Automated Testing Implementation
    doc.add_heading('3. Automated Testing Implementation', 1)
    
    doc.add_heading('3.1 Test Suite Architecture', 2)
    
    code_para = doc.add_paragraph()
    code_para.add_run('''import time
import json
import asyncio
from datetime import datetime
from collections import defaultdict
from typing import Dict, List, Any

class PerformanceTestSuite:
    def __init__(self):
        self.metrics = {
            "response_times": [],
            "intent_classification": {"correct": 0, "incorrect": 0},
            "citation_validity": [],
            "error_types": defaultdict(int),
            "token_usage": [],
            "cache_metrics": {"hits": 0, "misses": 0},
            "database_queries": []
        }
        self.test_results = []
    
    async def run_comprehensive_test(self):
        """Execute all test categories and collect metrics."""
        test_categories = [
            self.test_factual_accuracy,
            self.test_taxonomy_understanding,
            self.test_metadata_service,
            self.test_edge_cases,
            self.test_performance_limits
        ]
        
        for test_func in test_categories:
            await test_func()
        
        return self.generate_report()
    
    def generate_report(self) -> Dict[str, Any]:
        """Generate comprehensive performance report."""
        return {
            "timestamp": datetime.now().isoformat(),
            "total_tests": len(self.test_results),
            "passed": sum(1 for t in self.test_results if t["passed"]),
            "failed": sum(1 for t in self.test_results if not t["passed"]),
            "avg_response_time": sum(self.metrics["response_times"]) / len(self.metrics["response_times"]),
            "p95_response_time": self.calculate_percentile(95),
            "accuracy_rate": self.calculate_accuracy_rate(),
            "error_breakdown": dict(self.metrics["error_types"]),
            "cache_hit_rate": self.calculate_cache_hit_rate()
        }''').font.name = 'Courier New'
    
    doc.add_heading('3.2 Performance Monitoring Decorators', 2)
    
    code_para = doc.add_paragraph()
    code_para.add_run('''import functools
import tracemalloc
import psutil

def performance_monitor(func):
    """Decorator to monitor function performance."""
    @functools.wraps(func)
    async def wrapper(*args, **kwargs):
        # Start monitoring
        start_time = time.perf_counter()
        tracemalloc.start()
        start_memory = psutil.Process().memory_info().rss / 1024 / 1024  # MB
        
        try:
            result = await func(*args, **kwargs)
            success = True
            error = None
        except Exception as e:
            result = None
            success = False
            error = str(e)
        
        # End monitoring
        end_time = time.perf_counter()
        current, peak = tracemalloc.get_traced_memory()
        tracemalloc.stop()
        end_memory = psutil.Process().memory_info().rss / 1024 / 1024  # MB
        
        # Log metrics
        metrics = {
            "function": func.__name__,
            "execution_time": end_time - start_time,
            "memory_used": end_memory - start_memory,
            "peak_memory": peak / 1024 / 1024,  # MB
            "success": success,
            "error": error
        }
        
        log_performance_metrics(metrics)
        return result
    
    return wrapper''').font.name = 'Courier New'
    
    # 4. Load and Stress Testing
    doc.add_heading('4. Load and Stress Testing', 1)
    
    doc.add_heading('4.1 Concurrent User Simulation', 2)
    
    code_para = doc.add_paragraph()
    code_para.add_run('''async def simulate_concurrent_users(num_users: int, duration: int, queries_per_second: float):
    """Simulate concurrent users accessing the chatbot."""
    
    async def user_session(user_id: int):
        session_metrics = []
        query_interval = 1.0 / queries_per_second
        
        start_time = time.time()
        while time.time() - start_time < duration:
            query = random.choice(TEST_QUERIES)
            
            query_start = time.time()
            try:
                response = await send_query(query)
                latency = time.time() - query_start
                session_metrics.append({
                    "user_id": user_id,
                    "timestamp": time.time(),
                    "latency": latency,
                    "success": True
                })
            except Exception as e:
                session_metrics.append({
                    "user_id": user_id,
                    "timestamp": time.time(),
                    "error": str(e),
                    "success": False
                })
            
            await asyncio.sleep(query_interval)
        
        return session_metrics
    
    # Run concurrent sessions
    tasks = [user_session(i) for i in range(num_users)]
    results = await asyncio.gather(*tasks)
    
    return analyze_load_test_results(results)''').font.name = 'Courier New'
    
    doc.add_heading('4.2 Stress Test Scenarios', 2)
    
    # Add stress test table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Scenario'
    hdr_cells[1].text = 'Users'
    hdr_cells[2].text = 'Duration'
    hdr_cells[3].text = 'Expected Outcome'
    
    scenarios = [
        ('Baseline Load', '10', '5 min', 'All requests succeed, <2s latency'),
        ('Normal Load', '50', '10 min', '95% succeed, <3s P95 latency'),
        ('High Load', '100', '10 min', '90% succeed, <5s P95 latency'),
        ('Stress Test', '200', '5 min', 'Identify breaking point'),
        ('Spike Test', '10→100→10', '15 min', 'Graceful scaling'),
        ('Endurance Test', '50', '60 min', 'No memory leaks')
    ]
    
    for scenario in scenarios:
        row_cells = table.add_row().cells
        for i, value in enumerate(scenario):
            row_cells[i].text = value
    
    # 5. Database and Cache Performance
    doc.add_heading('5. Database and Cache Performance', 1)
    
    doc.add_heading('5.1 Query Optimization Testing', 2)
    
    code_para = doc.add_paragraph()
    code_para.add_run('''def test_database_performance():
    """Test ChromaDB query performance."""
    
    test_cases = [
        {
            "name": "Simple keyword search",
            "query": "privacy risks",
            "expected_time": 0.5
        },
        {
            "name": "Complex semantic search",
            "query": "What are the unintended consequences of AI in healthcare?",
            "expected_time": 1.0
        },
        {
            "name": "Metadata filtering",
            "query": "risks where domain='discrimination' AND timing='post-deployment'",
            "expected_time": 0.3
        }
    ]
    
    for test in test_cases:
        start = time.time()
        results = vectorstore.similarity_search(
            test["query"],
            k=10,
            filter=test.get("filter", None)
        )
        elapsed = time.time() - start
        
        assert elapsed < test["expected_time"], f"{test['name']} too slow: {elapsed}s"
        assert len(results) > 0, f"{test['name']} returned no results"''').font.name = 'Courier New'
    
    doc.add_heading('5.2 Cache Effectiveness Metrics', 2)
    
    metrics_list = [
        'Cache hit rate per query type',
        'Cache size growth over time',
        'Cache eviction rate',
        'Time saved by cache hits',
        'Memory usage by cache',
        'Cache warm-up time after restart'
    ]
    for metric in metrics_list:
        doc.add_paragraph(metric, style='List Bullet')
    
    # 6. Token Usage Optimization
    doc.add_heading('6. Token Usage Optimization', 1)
    
    doc.add_heading('6.1 Token Tracking Implementation', 2)
    
    code_para = doc.add_paragraph()
    code_para.add_run('''class TokenUsageTracker:
    def __init__(self):
        self.usage_history = []
        self.cost_per_1k_input = 0.01  # Example rate
        self.cost_per_1k_output = 0.03  # Example rate
    
    def track_query(self, query: str, response: str, model: str):
        usage = {
            "timestamp": datetime.now(),
            "model": model,
            "input_tokens": self.count_tokens(query),
            "output_tokens": self.count_tokens(response),
            "total_tokens": None,
            "estimated_cost": None
        }
        
        usage["total_tokens"] = usage["input_tokens"] + usage["output_tokens"]
        usage["estimated_cost"] = (
            (usage["input_tokens"] / 1000) * self.cost_per_1k_input +
            (usage["output_tokens"] / 1000) * self.cost_per_1k_output
        )
        
        self.usage_history.append(usage)
        return usage
    
    def get_optimization_recommendations(self):
        avg_input = sum(u["input_tokens"] for u in self.usage_history) / len(self.usage_history)
        avg_output = sum(u["output_tokens"] for u in self.usage_history) / len(self.usage_history)
        
        recommendations = []
        if avg_input > 500:
            recommendations.append("Consider prompt optimization to reduce input tokens")
        if avg_output > 1000:
            recommendations.append("Consider response length limits")
        
        return recommendations''').font.name = 'Courier New'
    
    # 7. Citation Validation
    doc.add_heading('7. Citation Validation Framework', 1)
    
    doc.add_heading('7.1 Automated Citation Testing', 2)
    
    code_para = doc.add_paragraph()
    code_para.add_run('''def validate_citations(response: str, documents: List[Document]) -> Dict:
    """Validate all citations in a response."""
    
    validation_results = {
        "total_citations": 0,
        "valid_citations": 0,
        "broken_links": [],
        "unsupported_claims": [],
        "citation_coverage": 0.0
    }
    
    # Extract all RID citations
    rid_pattern = r'RID-\\d{5}'
    citations = re.findall(rid_pattern, response)
    validation_results["total_citations"] = len(citations)
    
    # Validate each citation
    for rid in citations:
        # Check if RID exists in documents
        if any(doc.metadata.get("rid") == rid for doc in documents):
            validation_results["valid_citations"] += 1
            
            # Verify snippet is accessible
            snippet_path = f"/snippet/{rid}"
            if not check_snippet_exists(rid):
                validation_results["broken_links"].append(rid)
        else:
            validation_results["unsupported_claims"].append(rid)
    
    # Calculate citation coverage
    if validation_results["total_citations"] > 0:
        validation_results["citation_coverage"] = (
            validation_results["valid_citations"] / 
            validation_results["total_citations"]
        )
    
    return validation_results''').font.name = 'Courier New'
    
    # 8. Regression Testing
    doc.add_heading('8. Regression Testing Strategy', 1)
    
    doc.add_heading('8.1 Continuous Integration Pipeline', 2)
    
    code_para = doc.add_paragraph()
    code_para.add_run('''# .github/workflows/performance-tests.yml
name: Performance Regression Tests

on:
  push:
    branches: [main, develop]
  pull_request:
    branches: [main]

jobs:
  performance-test:
    runs-on: ubuntu-latest
    
    steps:
    - uses: actions/checkout@v2
    
    - name: Set up Python
      uses: actions/setup-python@v2
      with:
        python-version: '3.11'
    
    - name: Install dependencies
      run: |
        pip install -r requirements.txt
        pip install pytest-benchmark
    
    - name: Run performance tests
      run: |
        python -m pytest tests/performance/ \\
          --benchmark-only \\
          --benchmark-compare \\
          --benchmark-autosave
    
    - name: Check performance regression
      run: |
        python scripts/check_regression.py
    
    - name: Upload results
      uses: actions/upload-artifact@v2
      with:
        name: performance-results
        path: .benchmarks/''').font.name = 'Courier New'
    
    doc.add_heading('8.2 Performance Baseline Management', 2)
    
    doc.add_paragraph('Baseline metrics to track across releases:')
    
    baseline_metrics = [
        'Average response time for each query category',
        'P50, P95, P99 latency percentiles',
        'Memory usage under standard load',
        'Database query execution times',
        'Token usage per query type',
        'Error rates by category',
        'Cache effectiveness metrics'
    ]
    for metric in baseline_metrics:
        doc.add_paragraph(metric, style='List Bullet')
    
    # 9. Performance Monitoring Dashboard
    doc.add_heading('9. Performance Monitoring Dashboard', 1)
    
    doc.add_heading('9.1 Real-time Metrics Collection', 2)
    
    code_para = doc.add_paragraph()
    code_para.add_run('''from prometheus_client import Counter, Histogram, Gauge
import time

# Define metrics
query_counter = Counter('chatbot_queries_total', 'Total number of queries')
query_duration = Histogram('chatbot_query_duration_seconds', 'Query duration')
active_sessions = Gauge('chatbot_active_sessions', 'Number of active sessions')
error_counter = Counter('chatbot_errors_total', 'Total errors', ['error_type'])

@query_duration.time()
def process_query_with_metrics(query: str):
    """Process query with automatic metrics collection."""
    query_counter.inc()
    
    try:
        with active_sessions.track_inprogress():
            response = process_query(query)
        return response
    except Exception as e:
        error_counter.labels(error_type=type(e).__name__).inc()
        raise''').font.name = 'Courier New'
    
    doc.add_heading('9.2 Alert Thresholds', 2)
    
    # Add alert table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Warning Threshold'
    hdr_cells[2].text = 'Critical Threshold'
    
    alerts = [
        ('Response Time P95', '>3 seconds', '>5 seconds'),
        ('Error Rate', '>5%', '>10%'),
        ('Memory Usage', '>80%', '>90%'),
        ('Cache Hit Rate', '<60%', '<40%'),
        ('Active Sessions', '>150', '>200'),
        ('Token Usage/Query', '>3000', '>5000')
    ]
    
    for alert in alerts:
        row_cells = table.add_row().cells
        for i, value in enumerate(alert):
            row_cells[i].text = value
    
    # 10. Optimization Recommendations
    doc.add_heading('10. Performance Optimization Checklist', 1)
    
    doc.add_heading('10.1 Query Processing Optimizations', 2)
    optimizations = [
        'Implement query result caching with Redis',
        'Use connection pooling for database connections',
        'Optimize vector similarity search with FAISS indices',
        'Implement query batching for bulk operations',
        'Add request deduplication for identical concurrent queries'
    ]
    for opt in optimizations:
        doc.add_paragraph(opt, style='List Bullet')
    
    doc.add_heading('10.2 Response Generation Optimizations', 2)
    response_opts = [
        'Implement streaming responses for long outputs',
        'Use response templates for common queries',
        'Cache frequently used document snippets',
        'Optimize prompt engineering to reduce token usage',
        'Implement progressive disclosure for complex responses'
    ]
    for opt in response_opts:
        doc.add_paragraph(opt, style='List Bullet')
    
    # 11. Deliverables
    doc.add_heading('11. Testing Deliverables', 1)
    
    deliverables = [
        ('Week 1', 'Automated test suite implementation'),
        ('Week 1', 'Performance baseline report'),
        ('Week 2', 'Load test results and analysis'),
        ('Week 2', 'Database optimization recommendations'),
        ('Week 3', 'Token usage optimization report'),
        ('Week 3', 'CI/CD pipeline configuration'),
        ('Week 4', 'Performance monitoring dashboard'),
        ('Week 4', 'Final optimization recommendations')
    ]
    
    for week, deliverable in deliverables:
        p = doc.add_paragraph()
        p.add_run(f'{week}: ').bold = True
        p.add_run(deliverable)
    
    # Footer
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('MIT FutureTech - AI Risk Repository | Developer Testing Documentation').italic = True
    
    # Save document
    doc.save('DEVELOPER_PERFORMANCE_TESTING.docx')
    print('Created: DEVELOPER_PERFORMANCE_TESTING.docx')


def create_tester_evaluation_document():
    """Create comprehensive tester performance evaluation document."""
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = "Tester Performance Evaluation Guide"
    doc.core_properties.author = "MIT FutureTech"
    
    # Title
    title = doc.add_heading('AI Risk Repository Chatbot: Tester Performance Evaluation Guide', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Executive Summary
    doc.add_heading('Executive Summary', 1)
    doc.add_paragraph(
        'This document provides a comprehensive framework for external testers to evaluate the performance '
        'of the AI Risk Repository Chatbot. It includes structured testing protocols, evaluation rubrics, '
        'and detailed feedback collection instruments to assess both objective performance metrics and '
        'subjective user experience factors.'
    )
    
    # 1. Testing Overview
    doc.add_heading('1. Testing Overview', 1)
    
    doc.add_heading('1.1 What You Are Testing', 2)
    doc.add_paragraph(
        'The AI Risk Repository Chatbot is an intelligent assistant designed to help researchers, '
        'policymakers, and practitioners explore and understand AI-related risks. The system has access '
        'to a comprehensive database of 1,612 documented AI risks categorized across 7 domains.'
    )
    
    doc.add_heading('1.2 System Capabilities', 2)
    
    p = doc.add_paragraph()
    p.add_run('What the System CAN Do:').bold = True
    
    capabilities = [
        'Answer questions about AI risk categories and taxonomies',
        'Provide statistics and data about AI risks',
        'Cite academic sources with direct links to documents',
        'Explain how risks are organized and classified',
        'Search across multiple dimensions (domain, entity, timing, intent)',
        'Process queries in multiple languages',
        'Provide detailed explanations of specific risks'
    ]
    for cap in capabilities:
        doc.add_paragraph(f'✓ {cap}', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('What the System CANNOT Do:').bold = True
    
    limitations = [
        'Provide mitigation strategies or solutions',
        'Suggest governance frameworks',
        'Offer benchmarks or evaluation criteria',
        'Make predictions about future risks',
        'Give opinions or subjective assessments',
        'Access real-time or updated information',
        'Provide information outside the repository scope'
    ]
    for limit in limitations:
        doc.add_paragraph(f'✗ {limit}', style='List Bullet')
    
    # 2. Performance Evaluation Framework
    doc.add_heading('2. Performance Evaluation Framework', 1)
    
    doc.add_heading('2.1 Evaluation Dimensions', 2)
    
    # Add evaluation dimensions table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Dimension'
    hdr_cells[1].text = 'What We Measure'
    hdr_cells[2].text = 'How to Evaluate'
    
    dimensions = [
        ('Accuracy', 'Correctness of information provided', 'Verify facts against known data'),
        ('Relevance', 'How well responses match queries', 'Rate alignment with question intent'),
        ('Completeness', 'Thoroughness of responses', 'Check if all aspects addressed'),
        ('Clarity', 'Ease of understanding', 'Assess language and structure'),
        ('Speed', 'Response generation time', 'Measure time from query to response'),
        ('Reliability', 'Consistency across similar queries', 'Test variations of same question'),
        ('Error Handling', 'Grace in handling edge cases', 'Try ambiguous or invalid queries')
    ]
    
    for dim in dimensions:
        row_cells = table.add_row().cells
        for i, value in enumerate(dim):
            row_cells[i].text = value
    
    doc.add_heading('2.2 Performance Scoring Rubric', 2)
    
    doc.add_paragraph(
        'Use this 5-point scale to evaluate each dimension:'
    )
    
    rubric = [
        ('5 - Excellent', 'Exceeds expectations, professional quality'),
        ('4 - Good', 'Meets expectations, minor issues only'),
        ('3 - Satisfactory', 'Acceptable performance, some notable issues'),
        ('2 - Below Average', 'Significant issues affecting usability'),
        ('1 - Poor', 'Critical failures or unusable responses')
    ]
    
    for score, description in rubric:
        p = doc.add_paragraph()
        p.add_run(f'{score}: ').bold = True
        p.add_run(description)
    
    # 3. Structured Testing Protocol
    doc.add_heading('3. Structured Testing Protocol', 1)
    
    doc.add_heading('3.1 Test Session Structure', 2)
    doc.add_paragraph(
        'Each testing session should follow this structured approach to ensure comprehensive evaluation:'
    )
    
    session_structure = [
        ('0-5 minutes', 'Setup and familiarization'),
        ('5-20 minutes', 'Category A: Basic functionality testing'),
        ('20-35 minutes', 'Category B: Domain-specific testing'),
        ('35-50 minutes', 'Category C: Edge case testing'),
        ('50-65 minutes', 'Category D: Performance stress testing'),
        ('65-75 minutes', 'Category E: Comparative evaluation'),
        ('75-90 minutes', 'Documentation and feedback form completion')
    ]
    
    for time, activity in session_structure:
        p = doc.add_paragraph()
        p.add_run(f'{time}: ').bold = True
        p.add_run(activity)
    
    # 3.2 Category A: Basic Functionality
    doc.add_heading('3.2 Category A: Basic Functionality Testing', 2)
    
    doc.add_paragraph('Test the system\'s ability to handle fundamental queries:')
    
    doc.add_heading('Test Queries:', 3)
    basic_queries = [
        '"What types of AI risks exist?"',
        '"How many AI risks are in the repository?"',
        '"What are the 7 domains of AI risk?"',
        '"Tell me about privacy risks from AI"',
        '"How are AI risks categorized?"',
        '"What percentage of risks are caused by humans?"'
    ]
    
    for i, query in enumerate(basic_queries, 1):
        doc.add_paragraph(f'{i}. {query}')
    
    doc.add_heading('Evaluation Criteria:', 3)
    doc.add_paragraph('For each query, assess:')
    
    criteria = [
        'Response time (seconds): ________',
        'Accuracy (1-5): ⭐⭐⭐⭐⭐',
        'Completeness (1-5): ⭐⭐⭐⭐⭐',
        'Citations provided: Yes ☐ No ☐',
        'Citations work: Yes ☐ No ☐ N/A ☐',
        'Language appropriate: Yes ☐ No ☐',
        'Would a researcher find this useful: Yes ☐ No ☐'
    ]
    
    for criterion in criteria:
        doc.add_paragraph(criterion, style='List Bullet')
    
    # 3.3 Category B: Domain-Specific Testing
    doc.add_heading('3.3 Category B: Domain-Specific Testing', 2)
    
    doc.add_paragraph(
        'Test the system\'s depth of knowledge in specific domains:'
    )
    
    domain_tests = [
        ('Discrimination & Fairness', [
            '"What are the main discrimination risks in AI?"',
            '"How many discrimination risks are post-deployment?"',
            '"Show me unintentional discrimination risks"'
        ]),
        ('Privacy & Surveillance', [
            '"What privacy risks come from AI systems?"',
            '"Are privacy risks mostly intentional or unintentional?"',
            '"How do privacy risks relate to surveillance?"'
        ]),
        ('Misinformation', [
            '"What types of misinformation can AI create?"',
            '"Are misinformation risks pre or post deployment?"',
            '"Who causes misinformation risks - humans or AI?"'
        ])
    ]
    
    for domain, queries in domain_tests:
        doc.add_heading(f'{domain}:', 3)
        for query in queries:
            doc.add_paragraph(f'• {query}')
    
    # 3.4 Category C: Edge Case Testing
    doc.add_heading('3.4 Category C: Edge Case and Error Handling', 2)
    
    doc.add_paragraph(
        'Test how the system handles challenging or unusual inputs:'
    )
    
    doc.add_heading('Ambiguous Queries:', 3)
    ambiguous = [
        '"AI?"',
        '"Tell me about it"',
        '"What about risks?"',
        '"More"',
        '"Explain"'
    ]
    for query in ambiguous:
        doc.add_paragraph(f'• {query}')
    
    doc.add_heading('Out-of-Scope Queries:', 3)
    out_of_scope = [
        '"What are the best mitigation strategies?"',
        '"Which governance framework should I use?"',
        '"What will AI risks look like in 2030?"',
        '"How do I prevent these risks?"',
        '"What benchmarks measure AI safety?"'
    ]
    for query in out_of_scope:
        doc.add_paragraph(f'• {query}')
    
    doc.add_heading('Technical Stress Tests:', 3)
    stress = [
        'Very long query (500+ words)',
        'Query with special characters: !@#$%^&*()',
        'Query with typos: "Waht are teh mian rsiks?"',
        'Mixed language query',
        'Rapid sequential queries (5 in 10 seconds)'
    ]
    for test in stress:
        doc.add_paragraph(f'• {test}')
    
    # 3.5 Category D: Performance Stress Testing
    doc.add_heading('3.5 Category D: Performance Under Load', 2)
    
    doc.add_paragraph('Evaluate system performance under various conditions:')
    
    performance_tests = [
        ('Simple vs Complex Queries', 'Compare response times for basic vs elaborate questions'),
        ('Sequential Query Handling', 'Send 10 queries in rapid succession'),
        ('Context Retention', 'Ask follow-up questions to test conversation memory'),
        ('Multi-domain Queries', 'Ask questions spanning multiple risk domains'),
        ('Citation-heavy Responses', 'Request detailed information requiring many sources')
    ]
    
    for test_name, description in performance_tests:
        p = doc.add_paragraph()
        p.add_run(f'{test_name}: ').bold = True
        p.add_run(description)
    
    # 3.6 Category E: Comparative Evaluation
    doc.add_heading('3.6 Category E: Comparative Evaluation', 2)
    
    doc.add_paragraph(
        'Compare the chatbot\'s performance against alternative methods:'
    )
    
    doc.add_heading('Comparison Points:', 3)
    
    comparison_table = doc.add_table(rows=1, cols=4)
    comparison_table.style = 'Light Grid Accent 1'
    hdr_cells = comparison_table.rows[0].cells
    hdr_cells[0].text = 'Aspect'
    hdr_cells[1].text = 'This Chatbot'
    hdr_cells[2].text = 'ChatGPT/Claude'
    hdr_cells[3].text = 'Manual Search'
    
    aspects = [
        'Speed to answer',
        'Accuracy of facts',
        'Source citations',
        'Domain expertise',
        'Ease of use',
        'Trust in results'
    ]
    
    for aspect in aspects:
        row_cells = comparison_table.add_row().cells
        row_cells[0].text = aspect
        # Leave other cells empty for tester to fill
    
    # 4. Detailed Feedback Collection
    doc.add_heading('4. Detailed Feedback Collection Forms', 1)
    
    doc.add_heading('4.1 Quantitative Performance Metrics', 2)
    
    doc.add_paragraph('Please provide ratings for each performance aspect:')
    
    # Create performance rating table
    rating_table = doc.add_table(rows=1, cols=6)
    rating_table.style = 'Light Grid Accent 1'
    hdr_cells = rating_table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    for i in range(1, 6):
        hdr_cells[i].text = str(i)
    
    metrics = [
        'Response Accuracy',
        'Response Completeness',
        'Response Clarity',
        'Response Speed',
        'Citation Quality',
        'Error Handling',
        'Overall Usefulness',
        'System Reliability',
        'User Experience',
        'Trust in System'
    ]
    
    for metric in metrics:
        row_cells = rating_table.add_row().cells
        row_cells[0].text = metric
        # Leave rating cells empty for checkboxes
    
    doc.add_heading('4.2 Qualitative Feedback', 2)
    
    questions = [
        'What aspects of the chatbot\'s performance impressed you most?',
        'What aspects need the most improvement?',
        'Describe any responses that were particularly good or bad:',
        'What features or capabilities do you wish the chatbot had?',
        'How does this compare to your expectations?',
        'Would you recommend this tool to colleagues? Why or why not?',
        'What would make you trust the system\'s responses more?',
        'Describe any technical issues or bugs encountered:',
        'How could the response format be improved?',
        'Additional comments or suggestions:'
    ]
    
    for i, question in enumerate(questions, 1):
        doc.add_paragraph(f'{i}. {question}')
        doc.add_paragraph('_' * 60)
        doc.add_paragraph()  # Space for answer
    
    # 5. Performance Issue Classification
    doc.add_heading('5. Performance Issue Classification', 1)
    
    doc.add_heading('5.1 Issue Severity Levels', 2)
    
    severity_levels = [
        ('Critical', 'System failures, incorrect core information, security issues', 'Red'),
        ('Major', 'Significant inaccuracies, very slow responses, broken features', 'Orange'),
        ('Moderate', 'Minor inaccuracies, suboptimal responses, UI issues', 'Yellow'),
        ('Minor', 'Formatting issues, slight delays, preference matters', 'Green')
    ]
    
    for level, description, color in severity_levels:
        p = doc.add_paragraph()
        p.add_run(f'{level} ({color}): ').bold = True
        p.add_run(description)
    
    doc.add_heading('5.2 Issue Tracking Form', 2)
    
    doc.add_paragraph('Document each issue encountered:')
    
    issue_table = doc.add_table(rows=1, cols=5)
    issue_table.style = 'Light Grid Accent 1'
    hdr_cells = issue_table.rows[0].cells
    hdr_cells[0].text = 'Issue #'
    hdr_cells[1].text = 'Query/Context'
    hdr_cells[2].text = 'What Happened'
    hdr_cells[3].text = 'Expected Behavior'
    hdr_cells[4].text = 'Severity'
    
    # Add empty rows for issue documentation
    for i in range(1, 11):
        row_cells = issue_table.add_row().cells
        row_cells[0].text = str(i)
    
    # 6. Statistical Analysis Framework
    doc.add_heading('6. Statistical Analysis Framework', 1)
    
    doc.add_heading('6.1 Response Time Analysis', 2)
    
    doc.add_paragraph('Track response times across different query types:')
    
    time_table = doc.add_table(rows=1, cols=5)
    time_table.style = 'Light Grid Accent 1'
    hdr_cells = time_table.rows[0].cells
    hdr_cells[0].text = 'Query Type'
    hdr_cells[1].text = 'Min (s)'
    hdr_cells[2].text = 'Max (s)'
    hdr_cells[3].text = 'Avg (s)'
    hdr_cells[4].text = 'Sample Size'
    
    query_types = [
        'Simple factual',
        'Complex analytical',
        'Multi-domain',
        'Taxonomy queries',
        'Statistical queries',
        'Edge cases'
    ]
    
    for qtype in query_types:
        row_cells = time_table.add_row().cells
        row_cells[0].text = qtype
    
    doc.add_heading('6.2 Accuracy Assessment', 2)
    
    doc.add_paragraph('Evaluate accuracy across different information types:')
    
    accuracy_checks = [
        'Numerical statistics (counts, percentages)',
        'Domain names and categories',
        'Risk classifications (entity, timing, intent)',
        'Relationships between concepts',
        'Technical terminology usage',
        'Citation relevance and validity'
    ]
    
    for check in accuracy_checks:
        p = doc.add_paragraph()
        p.add_run(f'• {check}: ').bold = True
        p.add_run('Accurate ☐  Mostly Accurate ☐  Some Errors ☐  Many Errors ☐')
    
    # 7. User Journey Mapping
    doc.add_heading('7. User Journey Performance Mapping', 1)
    
    doc.add_heading('7.1 Typical Research Journey', 2)
    
    doc.add_paragraph(
        'Walk through a typical research scenario and evaluate performance at each step:'
    )
    
    journey_steps = [
        ('Initial Exploration', 'User asks broad question about AI risks'),
        ('Domain Discovery', 'User learns about the 7 risk domains'),
        ('Deep Dive', 'User focuses on specific domain of interest'),
        ('Data Gathering', 'User requests statistics and specifics'),
        ('Source Verification', 'User checks citations and sources'),
        ('Follow-up Questions', 'User asks clarifying questions'),
        ('Cross-domain Analysis', 'User explores relationships between domains')
    ]
    
    for step, description in journey_steps:
        p = doc.add_paragraph()
        p.add_run(f'{step}: ').bold = True
        p.add_run(description)
        doc.add_paragraph('Performance Rating: ⭐⭐⭐⭐⭐')
        doc.add_paragraph('Issues Encountered: _________________')
        doc.add_paragraph()
    
    doc.add_heading('7.2 Performance Bottlenecks', 2)
    
    doc.add_paragraph('Identify where users might encounter performance issues:')
    
    bottleneck_table = doc.add_table(rows=1, cols=4)
    bottleneck_table.style = 'Light Grid Accent 1'
    hdr_cells = bottleneck_table.rows[0].cells
    hdr_cells[0].text = 'Journey Stage'
    hdr_cells[1].text = 'Bottleneck Type'
    hdr_cells[2].text = 'Impact (1-5)'
    hdr_cells[3].text = 'Suggested Fix'
    
    # Add empty rows for bottleneck documentation
    for i in range(5):
        table.add_row()
    
    # 8. Comparative Analysis
    doc.add_heading('8. Comparative Performance Analysis', 1)
    
    doc.add_heading('8.1 Feature Comparison Matrix', 2)
    
    feature_table = doc.add_table(rows=1, cols=4)
    feature_table.style = 'Light Grid Accent 1'
    hdr_cells = feature_table.rows[0].cells
    hdr_cells[0].text = 'Feature'
    hdr_cells[1].text = 'AI Risk Chatbot'
    hdr_cells[2].text = 'General AI (GPT/Claude)'
    hdr_cells[3].text = 'Winner'
    
    features = [
        'Domain expertise',
        'Response speed',
        'Source citations',
        'Factual accuracy',
        'Query understanding',
        'Error handling',
        'Conversation flow',
        'Technical depth',
        'User friendliness',
        'Trust/Credibility'
    ]
    
    for feature in features:
        row_cells = feature_table.add_row().cells
        row_cells[0].text = feature
    
    doc.add_heading('8.2 Use Case Suitability', 2)
    
    doc.add_paragraph('Rate suitability for different use cases (1-5 scale):')
    
    use_cases = [
        'Academic research on AI risks',
        'Policy development and analysis',
        'General public education',
        'Risk assessment for AI projects',
        'Teaching and education',
        'Quick fact checking',
        'In-depth analysis',
        'Exploratory research'
    ]
    
    for use_case in use_cases:
        doc.add_paragraph(f'• {use_case}: ⭐⭐⭐⭐⭐')
    
    # 9. Final Assessment
    doc.add_heading('9. Final Performance Assessment', 1)
    
    doc.add_heading('9.1 Overall Performance Summary', 2)
    
    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = 'Light Grid Accent 1'
    hdr_cells = summary_table.rows[0].cells
    hdr_cells[0].text = 'Category'
    hdr_cells[1].text = 'Grade (A-F)'
    
    categories = [
        'Technical Performance',
        'Response Quality',
        'User Experience',
        'Reliability',
        'Domain Coverage',
        'Overall System'
    ]
    
    for category in categories:
        row_cells = summary_table.add_row().cells
        row_cells[0].text = category
    
    doc.add_heading('9.2 Recommendation', 2)
    
    doc.add_paragraph('Based on your evaluation, would you recommend this system for:')
    
    recommendations = [
        'Production use by researchers: Yes ☐ No ☐ With improvements ☐',
        'Public deployment: Yes ☐ No ☐ With improvements ☐',
        'Integration with other tools: Yes ☐ No ☐ With improvements ☐',
        'Replacement of manual search: Yes ☐ No ☐ Partially ☐'
    ]
    
    for rec in recommendations:
        doc.add_paragraph(rec)
    
    doc.add_heading('9.3 Priority Improvements', 2)
    
    doc.add_paragraph('List the top 5 improvements needed (in priority order):')
    
    for i in range(1, 6):
        doc.add_paragraph(f'{i}. ________________________________')
        doc.add_paragraph('   Impact: High ☐ Medium ☐ Low ☐')
        doc.add_paragraph('   Effort: High ☐ Medium ☐ Low ☐')
        doc.add_paragraph()
    
    # 10. Appendices
    doc.add_heading('10. Appendices', 1)
    
    doc.add_heading('Appendix A: Test Query Bank', 2)
    
    doc.add_paragraph('Additional queries for extended testing:')
    
    query_bank = {
        'Factual Queries': [
            '"How many risks are in each domain?"',
            '"What percentage of risks are AI-caused vs human-caused?"',
            '"Which domain has the most pre-deployment risks?"'
        ],
        'Analytical Queries': [
            '"How do discrimination risks differ from privacy risks?"',
            '"What patterns exist in post-deployment risks?"',
            '"Which combinations of entity and intent are most common?"'
        ],
        'Exploratory Queries': [
            '"Tell me about emerging AI risks"',
            '"What should policymakers know about AI risks?"',
            '"How has the understanding of AI risks evolved?"'
        ]
    }
    
    for category, queries in query_bank.items():
        doc.add_heading(category, 3)
        for query in queries:
            doc.add_paragraph(f'• {query}')
    
    doc.add_heading('Appendix B: Performance Benchmarks', 2)
    
    doc.add_paragraph('Target performance metrics for reference:')
    
    benchmarks = [
        ('Response Time', '<2 seconds for simple, <5 seconds for complex'),
        ('Accuracy', '>95% for factual information'),
        ('Citation Rate', '>80% of claims should have citations'),
        ('Error Rate', '<5% queries should result in errors'),
        ('User Satisfaction', '>4/5 average rating')
    ]
    
    for metric, target in benchmarks:
        p = doc.add_paragraph()
        p.add_run(f'{metric}: ').bold = True
        p.add_run(target)
    
    # Footer
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('MIT FutureTech - AI Risk Repository | Tester Evaluation Documentation').italic = True
    
    # Save document
    doc.save('TESTER_PERFORMANCE_EVALUATION.docx')
    print('Created: TESTER_PERFORMANCE_EVALUATION.docx')


if __name__ == '__main__':
    print("Creating performance evaluation documents...")
    create_developer_testing_document()
    create_tester_evaluation_document()
    print("\nSuccessfully created:")
    print("1. DEVELOPER_PERFORMANCE_TESTING.docx")
    print("2. TESTER_PERFORMANCE_EVALUATION.docx")
    print("\nThese documents are ready to upload to Google Drive.")