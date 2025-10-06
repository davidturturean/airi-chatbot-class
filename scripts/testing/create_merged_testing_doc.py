#!/usr/bin/env python3
"""
Create a comprehensive, color-coded testing document merging UI/UX, Developer, and Tester evaluation plans.
"""

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    import sys
    import subprocess
    print("python-docx not found. Installing...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE

# Define color scheme
COLOR_UI_UX = RGBColor(0, 0, 0)        # Black for UI/UX Testing
COLOR_DEV = RGBColor(64, 64, 64)       # Dark Gray for Developer Testing
COLOR_TESTER = RGBColor(128, 128, 128)  # Medium Gray for Tester Performance

def add_colored_paragraph(doc, text, color, bold=False, italic=False):
    """Add a paragraph with colored text."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = color
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return p

def add_colored_bullet(doc, text, color):
    """Add a bullet point with colored text."""
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.color.rgb = color
    return p

def create_comprehensive_testing_document():
    """Create the merged testing document with color coding."""
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = "Comprehensive Testing & Feedback Plan"
    doc.core_properties.author = "MIT FutureTech"
    
    # Title
    title = doc.add_heading('AI Risk Repository Chatbot: Comprehensive Testing & Feedback Plan', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Color Legend
    doc.add_heading('Document Color Coding', 1)
    add_colored_paragraph(doc, '■ Black Text: UI/UX User Feedback & Interface Testing', COLOR_UI_UX)
    add_colored_paragraph(doc, '■ Dark Gray Text: Developer Performance & Technical Testing', COLOR_DEV)
    add_colored_paragraph(doc, '■ Medium Gray Text: External Tester Performance Evaluation', COLOR_TESTER)
    
    # Executive Summary
    doc.add_heading('Executive Summary', 1)
    doc.add_paragraph(
        'This comprehensive plan integrates three critical evaluation streams for the AI Risk Repository '
        'Chatbot: user interface experience, technical performance metrics, and response quality assessment. '
        'Each colored section represents a different testing approach, enabling parallel execution while '
        'maintaining clear ownership and accountability.'
    )
    
    # 1. Unified Testing Framework
    doc.add_heading('1. Unified Testing Framework', 1)
    
    doc.add_heading('1.1 Testing Dimensions', 2)
    
    # UI/UX Dimensions
    add_colored_paragraph(doc, 'User Interface & Experience:', COLOR_UI_UX, bold=True)
    add_colored_bullet(doc, 'Entry point optimization (direct navigation vs context)', COLOR_UI_UX)
    add_colored_bullet(doc, 'Interface design patterns (embedded chat vs floating widget)', COLOR_UI_UX)
    add_colored_bullet(doc, 'Mobile responsiveness and cross-browser compatibility', COLOR_UI_UX)
    add_colored_bullet(doc, 'User journey from discovery to task completion', COLOR_UI_UX)
    add_colored_bullet(doc, 'Visual design and information architecture', COLOR_UI_UX)
    
    # Developer Testing Dimensions
    add_colored_paragraph(doc, 'Technical Performance:', COLOR_DEV, bold=True)
    add_colored_bullet(doc, 'Response latency (P50, P95, P99 percentiles)', COLOR_DEV)
    add_colored_bullet(doc, 'System accuracy and factual correctness', COLOR_DEV)
    add_colored_bullet(doc, 'Database query efficiency and optimization', COLOR_DEV)
    add_colored_bullet(doc, 'Token usage and cost optimization', COLOR_DEV)
    add_colored_bullet(doc, 'Load handling and concurrent user support', COLOR_DEV)
    add_colored_bullet(doc, 'Error rates and recovery mechanisms', COLOR_DEV)
    
    # Tester Evaluation Dimensions
    add_colored_paragraph(doc, 'Response Quality & Usability:', COLOR_TESTER, bold=True)
    add_colored_bullet(doc, 'Answer completeness and relevance', COLOR_TESTER)
    add_colored_bullet(doc, 'Language clarity and technical appropriateness', COLOR_TESTER)
    add_colored_bullet(doc, 'Citation quality and source reliability', COLOR_TESTER)
    add_colored_bullet(doc, 'Query understanding and intent recognition', COLOR_TESTER)
    add_colored_bullet(doc, 'Comparative performance vs alternatives', COLOR_TESTER)
    
    doc.add_heading('1.2 Success Metrics Overview', 2)
    
    # Create unified metrics table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Category'
    hdr_cells[1].text = 'Metric'
    hdr_cells[2].text = 'Target'
    hdr_cells[3].text = 'Testing Type'
    
    metrics = [
        ('User Experience', 'Task Completion Rate', '>80%', 'UI/UX'),
        ('User Experience', 'Time to First Interaction', '<10s', 'UI/UX'),
        ('User Experience', 'Mobile Usability Score', '>4/5', 'UI/UX'),
        ('Performance', 'Response Time', '<2s P95', 'Developer'),
        ('Performance', 'Accuracy Rate', '>95%', 'Developer'),
        ('Performance', 'Uptime', '>99.9%', 'Developer'),
        ('Quality', 'User Satisfaction', '>4/5', 'Tester'),
        ('Quality', 'Citation Relevance', '>90%', 'Tester'),
        ('Quality', 'Query Success Rate', '>75%', 'Tester')
    ]
    
    for category, metric, target, test_type in metrics:
        row_cells = table.add_row().cells
        row_cells[0].text = category
        row_cells[1].text = metric
        row_cells[2].text = target
        row_cells[3].text = test_type
    
    # 2. Testing Methodology
    doc.add_heading('2. Integrated Testing Methodology', 1)
    
    doc.add_heading('2.1 Phase 1: Setup & Preparation (Week 1)', 2)
    
    add_colored_paragraph(doc, 'UI/UX Preparation:', COLOR_UI_UX, bold=True)
    add_colored_bullet(doc, 'Recruit 10-12 diverse user participants', COLOR_UI_UX)
    add_colored_bullet(doc, 'Prepare user journey scenarios', COLOR_UI_UX)
    add_colored_bullet(doc, 'Set up screen recording and analytics tools', COLOR_UI_UX)
    add_colored_bullet(doc, 'Create task completion checklists', COLOR_UI_UX)
    
    add_colored_paragraph(doc, 'Developer Setup:', COLOR_DEV, bold=True)
    add_colored_bullet(doc, 'Create ground truth dataset (100+ test queries)', COLOR_DEV)
    add_colored_bullet(doc, 'Build automated test suite infrastructure', COLOR_DEV)
    add_colored_bullet(doc, 'Configure performance monitoring tools', COLOR_DEV)
    add_colored_bullet(doc, 'Establish baseline performance metrics', COLOR_DEV)
    
    add_colored_paragraph(doc, 'Tester Preparation:', COLOR_TESTER, bold=True)
    add_colored_bullet(doc, 'Develop evaluation rubrics and scoring sheets', COLOR_TESTER)
    add_colored_bullet(doc, 'Create test query categories and examples', COLOR_TESTER)
    add_colored_bullet(doc, 'Brief testers on system capabilities and limitations', COLOR_TESTER)
    add_colored_bullet(doc, 'Prepare feedback collection instruments', COLOR_TESTER)
    
    doc.add_heading('2.2 Phase 2: Testing Execution (Weeks 2-3)', 2)
    
    add_colored_paragraph(doc, 'UI/UX Testing Sessions:', COLOR_UI_UX, bold=True)
    add_colored_bullet(doc, '45-minute moderated user sessions', COLOR_UI_UX)
    add_colored_bullet(doc, 'Think-aloud protocol for interface navigation', COLOR_UI_UX)
    add_colored_bullet(doc, 'Task-based scenarios (finding information, exploring domains)', COLOR_UI_UX)
    add_colored_bullet(doc, 'Post-session interviews and surveys', COLOR_UI_UX)
    add_colored_bullet(doc, 'Heatmap and clickstream analysis', COLOR_UI_UX)
    
    add_colored_paragraph(doc, 'Developer Performance Testing:', COLOR_DEV, bold=True)
    add_colored_bullet(doc, 'Run automated test suite across all query categories', COLOR_DEV)
    add_colored_bullet(doc, 'Execute load testing (10, 50, 100, 200 concurrent users)', COLOR_DEV)
    add_colored_bullet(doc, 'Stress test with edge cases and malformed queries', COLOR_DEV)
    add_colored_bullet(doc, 'Monitor resource usage and optimization opportunities', COLOR_DEV)
    add_colored_bullet(doc, 'Validate citation accuracy and accessibility', COLOR_DEV)
    
    add_colored_paragraph(doc, 'Tester Quality Evaluation:', COLOR_TESTER, bold=True)
    add_colored_bullet(doc, 'Systematic testing of basic, complex, and edge queries', COLOR_TESTER)
    add_colored_bullet(doc, 'Domain-specific knowledge assessment', COLOR_TESTER)
    add_colored_bullet(doc, 'Response quality scoring on 5-point rubric', COLOR_TESTER)
    add_colored_bullet(doc, 'Comparative analysis with ChatGPT/Claude', COLOR_TESTER)
    add_colored_bullet(doc, 'Documentation of issues and improvement suggestions', COLOR_TESTER)
    
    doc.add_heading('2.3 Phase 3: Analysis & Synthesis (Week 4)', 2)
    
    doc.add_paragraph(
        'All three testing streams converge for integrated analysis:'
    )
    
    convergence_points = [
        'Cross-reference UI frustrations with technical bottlenecks',
        'Map user expectations to actual system capabilities',
        'Identify correlation between performance metrics and user satisfaction',
        'Prioritize improvements based on impact across all dimensions',
        'Create unified action plan with clear ownership'
    ]
    
    for point in convergence_points:
        doc.add_paragraph(point, style='List Bullet')
    
    # 3. Test Categories
    doc.add_heading('3. Comprehensive Test Categories', 1)
    
    doc.add_heading('3.1 Query Complexity Levels', 2)
    
    # Basic Queries
    doc.add_heading('Basic Queries (All Teams Test)', 3)
    queries = [
        '"What types of AI risks exist?"',
        '"How many AI risks are in the repository?"',
        '"Tell me about privacy risks"'
    ]
    for query in queries:
        doc.add_paragraph(f'• {query}')
    
    add_colored_paragraph(doc, 'UI/UX Focus: Ease of input, result presentation', COLOR_UI_UX, italic=True)
    add_colored_paragraph(doc, 'Developer Focus: Response time, accuracy', COLOR_DEV, italic=True)
    add_colored_paragraph(doc, 'Tester Focus: Completeness, clarity', COLOR_TESTER, italic=True)
    
    # Complex Queries
    doc.add_heading('Complex Analytical Queries', 3)
    queries = [
        '"Compare privacy risks to security risks"',
        '"What are the unintentional post-deployment discrimination risks?"',
        '"How do different domains intersect?"'
    ]
    for query in queries:
        doc.add_paragraph(f'• {query}')
    
    add_colored_paragraph(doc, 'UI/UX Focus: Information organization, navigation', COLOR_UI_UX, italic=True)
    add_colored_paragraph(doc, 'Developer Focus: Multi-service coordination, token usage', COLOR_DEV, italic=True)
    add_colored_paragraph(doc, 'Tester Focus: Analytical depth, coherence', COLOR_TESTER, italic=True)
    
    # Edge Cases
    doc.add_heading('Edge Cases & Error Scenarios', 3)
    edge_cases = [
        'Ambiguous queries: "AI?"',
        'Out-of-scope requests: "How to prevent risks?"',
        'Technical stress: Very long queries, special characters',
        'Language variations: Typos, non-English queries'
    ]
    for case in edge_cases:
        doc.add_paragraph(f'• {case}')
    
    add_colored_paragraph(doc, 'UI/UX Focus: Error messaging, graceful degradation', COLOR_UI_UX, italic=True)
    add_colored_paragraph(doc, 'Developer Focus: Error handling, system stability', COLOR_DEV, italic=True)
    add_colored_paragraph(doc, 'Tester Focus: User guidance, helpful alternatives', COLOR_TESTER, italic=True)
    
    # 4. Data Collection Framework
    doc.add_heading('4. Integrated Data Collection', 1)
    
    doc.add_heading('4.1 Quantitative Metrics', 2)
    
    # Create data collection table
    data_table = doc.add_table(rows=1, cols=3)
    data_table.style = 'Light Grid Accent 1'
    hdr_cells = data_table.rows[0].cells
    hdr_cells[0].text = 'Testing Type'
    hdr_cells[1].text = 'Metrics Collected'
    hdr_cells[2].text = 'Collection Method'
    
    data_points = [
        ('UI/UX', 'Click paths, time on task, abandonment points', 'Analytics tools, session recordings'),
        ('UI/UX', 'User satisfaction scores, preference rankings', 'Surveys, interviews'),
        ('Developer', 'Response times, error rates, resource usage', 'Automated monitoring, logs'),
        ('Developer', 'Accuracy scores, citation validity', 'Automated validation scripts'),
        ('Tester', 'Quality ratings, completeness scores', 'Evaluation rubrics'),
        ('Tester', 'Comparative rankings, trust metrics', 'Feedback forms')
    ]
    
    for test_type, metrics, method in data_points:
        row_cells = data_table.add_row().cells
        row_cells[0].text = test_type
        row_cells[1].text = metrics
        row_cells[2].text = method
    
    doc.add_heading('4.2 Qualitative Insights', 2)
    
    add_colored_paragraph(doc, 'UI/UX Qualitative Data:', COLOR_UI_UX, bold=True)
    add_colored_bullet(doc, 'User journey pain points and delights', COLOR_UI_UX)
    add_colored_bullet(doc, 'Mental models and expectations', COLOR_UI_UX)
    add_colored_bullet(doc, 'Feature requests and design preferences', COLOR_UI_UX)
    
    add_colored_paragraph(doc, 'Developer Observations:', COLOR_DEV, bold=True)
    add_colored_bullet(doc, 'Performance bottlenecks and optimization opportunities', COLOR_DEV)
    add_colored_bullet(doc, 'System behavior under stress conditions', COLOR_DEV)
    add_colored_bullet(doc, 'Resource utilization patterns', COLOR_DEV)
    
    add_colored_paragraph(doc, 'Tester Feedback:', COLOR_TESTER, bold=True)
    add_colored_bullet(doc, 'Response quality issues and patterns', COLOR_TESTER)
    add_colored_bullet(doc, 'Trust and credibility factors', COLOR_TESTER)
    add_colored_bullet(doc, 'Comparison with user expectations', COLOR_TESTER)
    
    # 5. Priority Matrix
    doc.add_heading('5. Unified Priority Matrix', 1)
    
    doc.add_heading('5.1 Issue Classification', 2)
    
    doc.add_paragraph('Issues are prioritized based on impact across all three testing dimensions:')
    
    # Priority table
    priority_table = doc.add_table(rows=1, cols=4)
    priority_table.style = 'Light Grid Accent 1'
    hdr_cells = priority_table.rows[0].cells
    hdr_cells[0].text = 'Priority'
    hdr_cells[1].text = 'Criteria'
    hdr_cells[2].text = 'Example Issues'
    hdr_cells[3].text = 'Timeline'
    
    priorities = [
        ('Critical', 'Affects all dimensions, blocks core functionality', 'System crashes, major inaccuracies, unusable interface', '48 hours'),
        ('High', 'Significant impact on 2+ dimensions', 'Slow responses affecting UX, poor mobile experience', '1 week'),
        ('Medium', 'Notable impact on 1 dimension', 'Suboptimal layouts, minor accuracy issues', '2 weeks'),
        ('Low', 'Minor improvements, nice-to-have', 'Formatting preferences, additional features', 'Backlog')
    ]
    
    for priority, criteria, examples, timeline in priorities:
        row_cells = priority_table.add_row().cells
        row_cells[0].text = priority
        row_cells[1].text = criteria
        row_cells[2].text = examples
        row_cells[3].text = timeline
    
    doc.add_heading('5.2 Cross-Functional Impact Analysis', 2)
    
    doc.add_paragraph(
        'Each identified issue is evaluated for its impact across all testing dimensions:'
    )
    
    # Impact analysis framework
    add_colored_paragraph(doc, 'Example: "Response takes 5+ seconds for complex queries"', COLOR_DEV, bold=True)
    add_colored_bullet(doc, 'UI/UX Impact: Users abandon tasks, poor experience', COLOR_UI_UX)
    add_colored_bullet(doc, 'Technical Impact: Performance SLA violation, resource strain', COLOR_DEV)
    add_colored_bullet(doc, 'Quality Impact: Perceived as unreliable, reduces trust', COLOR_TESTER)
    doc.add_paragraph('→ Classification: HIGH PRIORITY - affects all dimensions')
    
    # 6. Timeline and Resources
    doc.add_heading('6. Integrated Timeline & Resources', 1)
    
    doc.add_heading('6.1 Week-by-Week Breakdown', 2)
    
    # Timeline table
    timeline_table = doc.add_table(rows=1, cols=4)
    timeline_table.style = 'Light Grid Accent 1'
    hdr_cells = timeline_table.rows[0].cells
    hdr_cells[0].text = 'Week'
    hdr_cells[1].text = 'UI/UX Activities'
    hdr_cells[2].text = 'Developer Activities'
    hdr_cells[3].text = 'Tester Activities'
    
    weekly_activities = [
        ('Week 1', 'Recruit participants, prepare scenarios', 'Build test suite, create datasets', 'Develop rubrics, brief testers'),
        ('Week 2', 'Conduct user sessions (5)', 'Run automated tests, load testing', 'Begin systematic evaluation'),
        ('Week 3', 'Complete sessions, analyze data', 'Stress testing, optimization', 'Complete evaluations, comparative analysis'),
        ('Week 4', 'Synthesize findings', 'Performance report', 'Quality assessment report'),
        ('Week 5', 'Unified analysis and recommendations', 'All teams collaborate', 'All teams collaborate')
    ]
    
    for week, ui_ux, dev, tester in weekly_activities:
        row_cells = timeline_table.add_row().cells
        row_cells[0].text = week
        row_cells[1].text = ui_ux
        row_cells[2].text = dev
        row_cells[3].text = tester
    
    doc.add_heading('6.2 Resource Requirements', 2)
    
    add_colored_paragraph(doc, 'UI/UX Testing:', COLOR_UI_UX, bold=True)
    add_colored_bullet(doc, '10-12 user participants (45 min each)', COLOR_UI_UX)
    add_colored_bullet(doc, '2 UX researchers/facilitators', COLOR_UI_UX)
    add_colored_bullet(doc, 'Recording and analysis tools', COLOR_UI_UX)
    
    add_colored_paragraph(doc, 'Developer Testing:', COLOR_DEV, bold=True)
    add_colored_bullet(doc, '1-2 developers for test implementation', COLOR_DEV)
    add_colored_bullet(doc, 'CI/CD pipeline access', COLOR_DEV)
    add_colored_bullet(doc, 'Performance monitoring infrastructure', COLOR_DEV)
    
    add_colored_paragraph(doc, 'Tester Evaluation:', COLOR_TESTER, bold=True)
    add_colored_bullet(doc, '6-8 domain experts and researchers', COLOR_TESTER)
    add_colored_bullet(doc, '2 coordinators for session management', COLOR_TESTER)
    add_colored_bullet(doc, 'Evaluation tools and forms', COLOR_TESTER)
    
    # 7. Deliverables
    doc.add_heading('7. Integrated Deliverables', 1)
    
    doc.add_heading('7.1 Individual Reports', 2)
    
    add_colored_paragraph(doc, 'UI/UX Deliverables:', COLOR_UI_UX, bold=True)
    add_colored_bullet(doc, 'User journey maps with pain points', COLOR_UI_UX)
    add_colored_bullet(doc, 'Interface design recommendations', COLOR_UI_UX)
    add_colored_bullet(doc, 'Usability test results and videos', COLOR_UI_UX)
    
    add_colored_paragraph(doc, 'Developer Deliverables:', COLOR_DEV, bold=True)
    add_colored_bullet(doc, 'Performance baseline report', COLOR_DEV)
    add_colored_bullet(doc, 'Load test results and analysis', COLOR_DEV)
    add_colored_bullet(doc, 'Optimization recommendations', COLOR_DEV)
    
    add_colored_paragraph(doc, 'Tester Deliverables:', COLOR_TESTER, bold=True)
    add_colored_bullet(doc, 'Quality assessment scores', COLOR_TESTER)
    add_colored_bullet(doc, 'Comparative analysis report', COLOR_TESTER)
    add_colored_bullet(doc, 'Improvement suggestions', COLOR_TESTER)
    
    doc.add_heading('7.2 Unified Deliverable', 2)
    
    doc.add_paragraph(
        'Final Comprehensive Report including:'
    )
    
    final_deliverables = [
        'Executive summary with key findings across all dimensions',
        'Integrated priority matrix with cross-functional impacts',
        'Unified improvement roadmap with clear ownership',
        'Success metrics dashboard combining all KPIs',
        'Best practices and lessons learned',
        'Recommendations for continuous improvement'
    ]
    
    for deliverable in final_deliverables:
        doc.add_paragraph(deliverable, style='List Bullet')
    
    # 8. Success Criteria
    doc.add_heading('8. Definition of Success', 1)
    
    doc.add_heading('8.1 Minimum Viable Performance', 2)
    
    mvp_criteria = [
        ('User Experience', 'Can complete basic tasks without assistance', 'UI/UX'),
        ('User Experience', 'Mobile interface is functional', 'UI/UX'),
        ('Technical', 'Response time <5s for 95% of queries', 'Developer'),
        ('Technical', 'System uptime >99%', 'Developer'),
        ('Quality', 'Accuracy >85% for factual queries', 'Tester'),
        ('Quality', 'Users rate experience >3/5', 'Tester')
    ]
    
    for category, criterion, test_type in mvp_criteria:
        p = doc.add_paragraph()
        p.add_run(f'• {category}: ').bold = True
        p.add_run(criterion)
        if test_type == 'UI/UX':
            p.runs[1].font.color.rgb = COLOR_UI_UX
        elif test_type == 'Developer':
            p.runs[1].font.color.rgb = COLOR_DEV
        else:
            p.runs[1].font.color.rgb = COLOR_TESTER
    
    doc.add_heading('8.2 Target Performance', 2)
    
    target_criteria = [
        ('User Experience', 'Intuitive interface requiring no training', 'UI/UX'),
        ('User Experience', 'Seamless mobile experience', 'UI/UX'),
        ('Technical', 'Response time <2s for 95% of queries', 'Developer'),
        ('Technical', 'System uptime >99.9%', 'Developer'),
        ('Quality', 'Accuracy >95% for all queries', 'Tester'),
        ('Quality', 'Users rate experience >4.5/5', 'Tester')
    ]
    
    for category, criterion, test_type in target_criteria:
        p = doc.add_paragraph()
        p.add_run(f'• {category}: ').bold = True
        p.add_run(criterion)
        if test_type == 'UI/UX':
            p.runs[1].font.color.rgb = COLOR_UI_UX
        elif test_type == 'Developer':
            p.runs[1].font.color.rgb = COLOR_DEV
        else:
            p.runs[1].font.color.rgb = COLOR_TESTER
    
    # 9. Risk Mitigation
    doc.add_heading('9. Risk Mitigation Strategy', 1)
    
    risks = [
        ('Low participant engagement', 'Recruit 20% extra participants, offer incentives', 'UI/UX'),
        ('Test environment instability', 'Use staging environment, have rollback plan', 'Developer'),
        ('Tester bias', 'Use diverse tester pool, structured rubrics', 'Tester'),
        ('Timeline slippage', 'Build in buffer time, parallel execution where possible', 'All'),
        ('Conflicting feedback', 'Use priority matrix, data-driven decisions', 'All')
    ]
    
    for risk, mitigation, owner in risks:
        p = doc.add_paragraph()
        p.add_run(f'• Risk: ').bold = True
        p.add_run(f'{risk} → ')
        p.add_run(f'Mitigation: {mitigation}')
        if owner == 'UI/UX':
            p.runs[2].font.color.rgb = COLOR_UI_UX
        elif owner == 'Developer':
            p.runs[2].font.color.rgb = COLOR_DEV
        elif owner == 'Tester':
            p.runs[2].font.color.rgb = COLOR_TESTER
    
    # 10. Next Steps
    doc.add_heading('10. Immediate Next Steps', 1)
    
    next_steps = [
        ('Day 1-2', 'Approve this comprehensive plan', 'All Teams'),
        ('Day 3-5', 'Begin participant recruitment', 'UI/UX Team'),
        ('Day 3-5', 'Set up test infrastructure', 'Developer Team'),
        ('Day 3-5', 'Finalize evaluation rubrics', 'Tester Team'),
        ('Day 6-7', 'Conduct pilot tests', 'All Teams'),
        ('Week 2', 'Begin full testing execution', 'All Teams')
    ]
    
    for timing, action, owner in next_steps:
        p = doc.add_paragraph()
        p.add_run(f'{timing}: ').bold = True
        p.add_run(f'{action} ')
        p.add_run(f'[{owner}]').italic = True
    
    # Questions for Discussion
    doc.add_heading('Questions for Stakeholder Discussion', 1)
    
    questions = [
        'What is the relative priority between UI/UX, performance, and quality?',
        'Are there specific user groups or use cases to prioritize?',
        'What is the budget allocation across the three testing streams?',
        'Should we coordinate with the main airisk.mit.edu site team?',
        'What level of performance degradation is acceptable during peak usage?',
        'How should we handle conflicting recommendations across teams?'
    ]
    
    for i, question in enumerate(questions, 1):
        doc.add_paragraph(f'{i}. {question}')
    
    # Footer
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('MIT FutureTech - AI Risk Repository').italic = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Comprehensive Testing & Feedback Plan | November 2024').italic = True
    
    # Save document
    doc.save('COMPREHENSIVE_TESTING_PLAN.docx')
    print('Successfully created: COMPREHENSIVE_TESTING_PLAN.docx')
    print('Color coding:')
    print('  - Black: UI/UX Testing')
    print('  - Dark Gray: Developer Performance Testing')
    print('  - Medium Gray: Tester Performance Evaluation')


if __name__ == '__main__':
    create_comprehensive_testing_document()