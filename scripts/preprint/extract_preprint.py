#!/usr/bin/env python3
"""
Extract text content from the AI Risk Repository Preprint DOCX file.
Preserves structure including headings, paragraphs, tables, and lists.
"""

import os
from pathlib import Path
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
import re

def extract_text_from_docx(docx_path):
    """Extract structured text from a DOCX file."""
    
    doc = Document(docx_path)
    extracted_content = []
    current_section = None
    section_counter = 0
    
    print(f"Processing document with {len(doc.paragraphs)} paragraphs...")
    
    for element in doc.element.body:
        # Check if element is a paragraph
        if element.tag.endswith('p'):
            para = Paragraph(element, doc)
            text = para.text.strip()
            
            if not text:
                continue
                
            # Check if it's a heading based on style
            if para.style and para.style.name:
                style_name = para.style.name.lower()
                
                if 'heading' in style_name:
                    # Extract heading level
                    level = 1
                    if 'heading 1' in style_name:
                        level = 1
                        section_counter += 1
                        current_section = text
                        extracted_content.append(f"\n\n{'=' * 80}")
                        extracted_content.append(f"SECTION {section_counter}: {text}")
                        extracted_content.append('=' * 80)
                    elif 'heading 2' in style_name:
                        level = 2
                        extracted_content.append(f"\n\n## {text}")
                        extracted_content.append('-' * 40)
                    elif 'heading 3' in style_name:
                        level = 3
                        extracted_content.append(f"\n### {text}")
                    else:
                        extracted_content.append(f"\n#### {text}")
                else:
                    # Regular paragraph
                    extracted_content.append(f"\n{text}")
            else:
                # No style, treat as regular paragraph
                extracted_content.append(f"\n{text}")
                
        # Check if element is a table
        elif element.tag.endswith('tbl'):
            table = Table(element, doc)
            extracted_content.append("\n\n[TABLE]")
            
            # Extract table content
            for i, row in enumerate(table.rows):
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip().replace('\n', ' ')
                    row_text.append(cell_text)
                
                if i == 0:
                    # Header row
                    extracted_content.append("| " + " | ".join(row_text) + " |")
                    extracted_content.append("|" + "---|" * len(row_text))
                else:
                    extracted_content.append("| " + " | ".join(row_text) + " |")
            
            extracted_content.append("[/TABLE]\n")
    
    # Join all content
    full_text = ''.join(extracted_content)
    
    # Clean up excessive whitespace
    full_text = re.sub(r'\n{4,}', '\n\n\n', full_text)
    
    return full_text, section_counter

def extract_metadata(doc_path):
    """Extract document metadata."""
    doc = Document(doc_path)
    
    metadata = {
        'total_paragraphs': len(doc.paragraphs),
        'total_tables': len(doc.tables),
        'core_properties': {}
    }
    
    # Extract core properties if available
    try:
        core_props = doc.core_properties
        metadata['core_properties'] = {
            'title': core_props.title or 'AI Risk Repository Preprint',
            'author': core_props.author or 'Unknown',
            'created': str(core_props.created) if core_props.created else None,
            'modified': str(core_props.modified) if core_props.modified else None,
            'subject': core_props.subject or None,
            'keywords': core_props.keywords or None
        }
    except:
        pass
    
    return metadata

def main():
    # Path to the DOCX file
    docx_path = Path('/Users/davidturturean/Documents/Codingprojects/airi-chatbot-class/data/info_files/AI_Risk_Repository_Preprint.docx')
    
    if not docx_path.exists():
        print(f"Error: File not found at {docx_path}")
        return
    
    print(f"Extracting text from: {docx_path}")
    print(f"File size: {docx_path.stat().st_size / 1024 / 1024:.2f} MB")
    
    # Extract metadata
    print("\nExtracting metadata...")
    metadata = extract_metadata(docx_path)
    print(f"Found {metadata['total_paragraphs']} paragraphs and {metadata['total_tables']} tables")
    
    # Extract text content
    print("\nExtracting text content...")
    text_content, num_sections = extract_text_from_docx(docx_path)
    
    # Save to file
    output_path = Path('/Users/davidturturean/Documents/Codingprojects/airi-chatbot-class/data/info_files/preprint_full.txt')
    
    print(f"\nWriting to: {output_path}")
    with open(output_path, 'w', encoding='utf-8') as f:
        # Write metadata header
        f.write("AI RISK REPOSITORY PREPRINT - FULL TEXT EXTRACTION\n")
        f.write("=" * 80 + "\n")
        f.write(f"Document: {metadata['core_properties'].get('title', 'Unknown')}\n")
        f.write(f"Total Sections: {num_sections}\n")
        f.write(f"Total Paragraphs: {metadata['total_paragraphs']}\n")
        f.write(f"Total Tables: {metadata['total_tables']}\n")
        f.write("=" * 80 + "\n\n")
        
        # Write content
        f.write(text_content)
    
    # Calculate statistics
    word_count = len(text_content.split())
    char_count = len(text_content)
    
    print(f"\nExtraction complete!")
    print(f"- Sections found: {num_sections}")
    print(f"- Words extracted: {word_count:,}")
    print(f"- Characters extracted: {char_count:,}")
    print(f"- Output file size: {output_path.stat().st_size / 1024:.2f} KB")
    
    # Also save a raw version without formatting
    raw_output_path = Path('/Users/davidturturean/Documents/Codingprojects/airi-chatbot-class/data/info_files/preprint_raw.txt')
    raw_text = re.sub(r'[=\-]{3,}', '', text_content)  # Remove section dividers
    raw_text = re.sub(r'SECTION \d+:', '', raw_text)  # Remove section markers
    raw_text = re.sub(r'\[TABLE\]|\[\/TABLE\]', '', raw_text)  # Remove table markers
    raw_text = re.sub(r'\n{3,}', '\n\n', raw_text)  # Normalize whitespace
    
    with open(raw_output_path, 'w', encoding='utf-8') as f:
        f.write(raw_text.strip())
    
    print(f"- Raw text saved to: {raw_output_path}")

if __name__ == "__main__":
    main()